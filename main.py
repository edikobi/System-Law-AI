"""
Единая точка входа для системы ИИ-юриста.
Режимы: интерактивный, индексация, Telegram-бот.
"""

# Standard library
import sys
import os
import asyncio
import argparse
import json
import time
from pathlib import Path
from datetime import datetime
import re
from src.utils.encoding_detector import read_file_with_encoding
from config import get_law_display_name
import gc
import platform

# Third-party
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Project imports
from config import Config, LEGAL_CATEGORIES
from api_manager import api_manager
from src.classification.gigachat_classifier import GigaChatClassifier
from src.core.sequential_searcher import SequentialArticleSearcher
from src.rag.smart_rag_simple import SmartRAGSystem
from src.core.law_navigator import LawNavigator


PROJECT_ROOT = Path(__file__).parent
CHAPTER_DESCRIPTIONS_FILE = PROJECT_ROOT / "chapter_descriptions.json"


class LegalAISystem:
    """Главный класс системы ИИ-юриста, объединяющий классификацию, поиск и генерацию"""
    
    def __init__(self, auto_init_rag: bool = True):
        """Инициализация системы ИИ-юриста"""
        self.project_root = PROJECT_ROOT
        self.config = Config()
        self.classifier = GigaChatClassifier()
        self.rag_system = None
        self.navigator = None
        self.searcher = None
        self.exporter = DocxExporter(self.project_root)
        
        if auto_init_rag:
            self._ensure_rag_initialized()
    
    def _ensure_rag_initialized(self) -> None:
        """Инициализирует RAG систему если она еще не инициализирована"""
        if self.rag_system is not None:
            return
    
        try:
            self._check_and_rebuild_index()
            self.rag_system = SmartRAGSystem()
            # Используем rag_system как навигатор, так как он имеет метод find_relevant_containers()
            self.navigator = self.rag_system
            self.searcher = SequentialArticleSearcher(
                navigator=self.rag_system,  # SmartRAGSystem имеет find_relevant_containers()
                project_root=self.project_root
            )
        except Exception as e:
            print(f"❌ Ошибка инициализации RAG: {e}")
            raise
    
    def _check_and_rebuild_index(self) -> bool:
        """Проверяет и перестраивает индекс если необходимо"""
        chroma_db_path = self.project_root / "chroma_db"
        descriptions_path = CHAPTER_DESCRIPTIONS_FILE

        try:
            # Создаем временный экземпляр RAG без автоиндексации
            temp_rag = SmartRAGSystem(auto_index=False)

            # Проверяем три условия
            condition_1 = temp_rag.collection.count() > 0
            condition_2 = descriptions_path.exists()
            condition_3 = not temp_rag._check_law_files_changes()

            # Если все условия выполнены - переиндексация не требуется
            if condition_1 and condition_2 and condition_3:
                return False

        except Exception as e:
            print(f"⚠️ Ошибка при проверке состояния базы: {e}")
            # При ошибке проверки переходим к переиндексации

        # Требуется переиндексация
        print("🔄 Векторная база отсутствует, повреждена или обнаружены изменения в законах. Запускаю полную индексацию...")
        try:
            asyncio.run(self.index_laws_with_ai())
            # КРИТИЧНО: Создать новый экземпляр и СОХРАНИТЬ его
            self.rag_system = SmartRAGSystem(force_reindex=True)
            return True
        except Exception as e:
            print(f"❌ Ошибка переиндексации: {e}")
            return True
    
    async def process_question(self, question: str) -> dict:
        """Обрабатывает вопрос пользователя"""
        try:
            self._ensure_rag_initialized()
            
            print("🔍 Классификация вопроса...")
            category = self.classifier.classify_with_rag_assistance(question)
            print(f"📂 Категория: {category}")
            
            result = await self.searcher.sequential_search(question, category)
            result["category"] = category
            
            return result
        except Exception as e:
            return {
                "status": "error",
                "error": str(e)
            }
    
    async def index_laws_with_ai(self) -> dict:
        """Запускает AI-индексацию законов с созданием структурированного индекса"""
        print("🤖 Запуск AI-индексации законов...")

        from src.core.docx_intelligent_parser import DocxIntelligentParser

        parser = DocxIntelligentParser()
        laws_dir = self.project_root / "laws"

        if not laws_dir.exists():
            print(f"❌ Папка законов не найдена: {laws_dir}")
            return {"status": "error", "message": "Laws directory not found"}

        # Структура: {law_name: {"chapters": [{"title": ..., "description": ...}]}}
        law_index = {}
        total_chapters = 0
        total_api_calls = 0

        # Получаем список файлов законов
        law_files = list(laws_dir.glob("*.docx")) + list(laws_dir.glob("*.txt"))
        law_files = [f for f in law_files if not f.name.startswith('~$')]

        print(f"📚 Найдено законов: {len(law_files)}")

        for law_file in law_files:
            try:
                print(f"\n{'='*60}")
                print(f"📄 Обработка: {law_file.name}")
                print(f"{'='*60}")
        
                law_index[law_file.name] = {"chapters": []}
        
                if law_file.suffix.lower() == '.docx':
                    # ШАГ 1: Парсим DOCX → получаем дерево
                    root = parser.parse(str(law_file))
            
                    # ШАГ 2: Получаем чанки первого уровня (главы)
                    # Используем большой лимит, чтобы получить крупные главы
                    top_level_chunks = parser.chunk_tree(root, max_chars=999999999)
            
                    print(f"   📊 Найдено глав верхнего уровня: {len(top_level_chunks)}")
            
                    for i, chapter_chunk in enumerate(top_level_chunks, 1):
                        chapter_title = chapter_chunk.get('chunk_title', f'Раздел {i}')
                        chapter_content = chapter_chunk.get('content', '')
                
                        if not chapter_content.strip():
                            continue
                
                        print(f"\n   🔍 Глава {i}/{len(top_level_chunks)}: {chapter_title[:60]}...")
                        print(f"      Размер: {len(chapter_content)} символов")
                
                        # ШАГ 3: Проверяем размер главы
                        if len(chapter_content) > 100000:
                            print(f"      ⚠️  Глава слишком большая, разбиваем на подчанки...")

                            # Получаем ссылку на узел дерева для иерархического разбиения
                            source_node = chapter_chunk.get('source_node_ref')

                            if source_node:
                                # ИЕРАРХИЧЕСКОЕ разбиение: используем структуру документа
                                sub_chunks = parser.chunk_tree(source_node, max_chars=100000)
                                print(f"      📦 Использовано иерархическое разбиение: {len(sub_chunks)} подчанков")
                            else:
                                # FALLBACK: простое разбиение по символам (для старых файлов без ссылки)
                                sub_chunks = []
                                for j in range(0, len(chapter_content), 100000):
                                    sub_chunk_text = chapter_content[j:j+100000]
                                    sub_chunks.append({
                                    'chunk_title': f"{chapter_title} (часть {j//100000 + 1})",
                                    'content': sub_chunk_text
                                    })
                                    print(f"      📦 Использовано простое разбиение: {len(sub_chunks)} частей")
                    
                            # Анализируем каждый подчанк через ИИ
                            sub_descriptions = []
                            for k, sub_chunk in enumerate(sub_chunks, 1):
                                sub_title = sub_chunk.get('chunk_title', f"{chapter_title} (часть {k})")
                                sub_content = sub_chunk.get('content', '')
                        
                                print(f"         Анализ подчанка {k}/{len(sub_chunks)}: {sub_title[:40]}...")
                        
                                # ИИ читает ВЕСЬ текст подчанка
                                description = await self._generate_chapter_description(
                                    law_file.name,
                                    sub_title,
                                    sub_content  # ВЕСЬ текст подчанка, не обрезанный
                                )
                                sub_descriptions.append(description)
                                total_api_calls += 1
                    
                            # Объединяем описания подчанков в единое описание главы
                            combined_description = await self._combine_sub_descriptions(
                                law_file.name,
                                chapter_title,
                                sub_descriptions
                            )
                            total_api_calls += 1
                    
                            final_description = combined_description
                        else:
                            # Глава помещается целиком - анализируем весь текст
                            print(f"      ✅ Анализ главы целиком...")
                    
                            # ИИ читает ВЕСЬ текст главы
                            final_description = await self._generate_chapter_description(
                                law_file.name,
                                chapter_title,
                                chapter_content  # ВЕСЬ текст главы
                            )
                            total_api_calls += 1
                
                        # Сохраняем в индекс
                        law_index[law_file.name]["chapters"].append({
                            "title": chapter_title,
                            "description": final_description,
                            "size": len(chapter_content),
                            "level": chapter_chunk.get('level', 0)
                        })
                        total_chapters += 1
                
                        print(f"      ✅ Описание сохранено")
        
                else:
                    # Для TXT файлов используем простое чтение
                    print(f"   📄 TXT файл - читаем целиком")
                    content = read_file_with_encoding(law_file)
            
                    # Простое разбиение по главам для TXT
                    lines = content.split('\n')
                    chapter_pattern = re.compile(r'^(ГЛАВА|Глава|РАЗДЕЛ|Раздел)\s+', re.IGNORECASE)
            
                    current_chapter = []
                    current_title = f"{law_file.name} - Преамбула"
            
                    for line in lines:
                        if chapter_pattern.match(line.strip()):
                            # Сохраняем предыдущую главу
                            if current_chapter:
                                chapter_content = '\n'.join(current_chapter)
                                if len(chapter_content.strip()) > 100:
                                    description = await self._generate_chapter_description(
                                        law_file.name,
                                        current_title,
                                        chapter_content
                                    )
                                    law_index[law_file.name]["chapters"].append({
                                        "title": current_title,
                                        "description": description,
                                        "size": len(chapter_content),
                                        "level": 0
                                    })
                                    total_chapters += 1
                                    total_api_calls += 1
                    
                            # Начинаем новую главу
                            current_title = line.strip()
                            current_chapter = [line]
                        else:
                            current_chapter.append(line)
            
                    # Сохраняем последнюю главу
                    if current_chapter:
                        chapter_content = '\n'.join(current_chapter)
                        if len(chapter_content.strip()) > 100:
                            description = await self._generate_chapter_description(
                                law_file.name,
                                current_title,
                                chapter_content
                            )
                            law_index[law_file.name]["chapters"].append({
                                "title": current_title,
                                "description": description,
                                "size": len(chapter_content),
                                "level": 0
                            })
                            total_chapters += 1
                            total_api_calls += 1
    
            except Exception as e:
                print(f"⚠️ Ошибка обработки {law_file.name}: {e}")
                import traceback
                traceback.print_exc()
                continue

        # Сохраняем индекс
        index_file = self.project_root / "law_index.json"
        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(law_index, f, ensure_ascii=False, indent=2)

        print(f"\n{'='*60}")
        print(f"✅ Индекс сохранен: {index_file}")
        print(f"📊 Статистика:")
        print(f"   • Законов: {len(law_index)}")
        print(f"   • Глав: {total_chapters}")
        print(f"   • API вызовов: {total_api_calls}")
        print(f"{'='*60}")

        # Удаляем старую ChromaDB и создаем новую из индекса
        print("\n🔄 Подготовка к созданию векторной базы...")
    
        # Сначала закрываем существующую RAG систему если есть
        if self.rag_system is not None:
            print("🔄 Закрываю существующую RAG систему...")
            self.rag_system.clear_database()
            self.rag_system = None
            # Даём время на освобождение ресурсов
            gc.collect()
            if platform.system() == 'Windows':
                time.sleep(2.0)
    
        # Проверяем, осталась ли папка после очистки
        chroma_path = self.project_root / "chroma_db"
        if chroma_path.exists():
            print("🗑️ Удаляю остатки старой ChromaDB...")
            gc.collect()
            base_delay = 1.0 if platform.system() == 'Windows' else 0.3
        
            for attempt in range(5):
                try:
                    shutil.rmtree(chroma_path)
                    print("✅ Старая ChromaDB удалена")
                    break
                except (PermissionError, OSError) as e:
                    if attempt < 4:
                        gc.collect()
                        delay = base_delay * (attempt + 1)
                        print(f"⏳ Попытка {attempt + 1}/5 удаления ChromaDB... (ожидание {delay:.1f}с)")
                        time.sleep(delay)
                    else:
                        print(f"⚠️ Не удалось удалить старую ChromaDB: {e}")
                        print("   Продолжаю с существующей базой...")
    
        # Создаем новую ChromaDB из индекса
        print("\n🔄 Создание векторной базы из индекса...")
        self.rag_system = SmartRAGSystem()
    
        return {
            "status": "success",
            "indexed_laws": len(law_index),
            "indexed_chapters": total_chapters,
            "api_calls": total_api_calls
        }
    
    async def _generate_chapter_description(self, law_name: str, chapter_title: str, content: str) -> str:
        """Генерирует описание главы через DeepSeek API"""
        prompt = f"""Ты - юридический аналитик. Проанализируй следующий фрагмент закона и напиши КРАТКОЕ описание (2-3 предложения) о чём эта глава с юридической точки зрения.

    ЗАКОН: {law_name}
    ГЛАВА: {chapter_title}

    ТЕКСТ ГЛАВЫ (полностью):
    {content}

    ИНСТРУКЦИИ:
    1. Опиши основную тему главы
    2. Укажи ключевые правовые институты, которые регулируются
    3. Укажи, для каких ситуаций эта глава может быть полезна

    ФОРМАТ ОТВЕТА: Только текст описания, без заголовков и маркеров."""

        try:
            messages = [{"role": "user", "content": prompt}]
            response = api_manager.deepseek_completion(
                messages=messages,
                temperature=0.3
            )
            return response.strip()
        except Exception as e:
            print(f"⚠️ Ошибка генерации описания: {e}")
            return f"Раздел '{chapter_title}' закона '{law_name}'"

    async def _combine_sub_descriptions(self, law_name: str, chapter_title: str, sub_descriptions: list) -> str:
        """Объединяет описания подчанков в единое описание главы через DeepSeek"""
        prompt = f"""Ты - юридический аналитик. Перед тобой несколько описаний частей одной главы закона. Объедини их в единое связное описание (2-3 предложения).

    ЗАКОН: {law_name}
    ГЛАВА: {chapter_title}

    ОПИСАНИЯ ЧАСТЕЙ:
    {chr(10).join([f"{i+1}. {desc}" for i, desc in enumerate(sub_descriptions)])}

    ИНСТРУКЦИИ:
    1. Объедини ключевые темы из всех частей
    2. Сохрани юридическую точность
    3. Укажи общую область применения

    ФОРМАТ ОТВЕТА: Только текст описания, без заголовков и маркеров."""

        try:
            messages = [{"role": "user", "content": prompt}]
            response = api_manager.deepseek_completion(
                messages=messages,
                temperature=0.3
            )
            return response.strip()
        except Exception as e:
            print(f"⚠️ Ошибка объединения описаний: {e}")
            return " | ".join(sub_descriptions[:3])
    
    def export_to_docx(self, result: dict, output_path: Path = None) -> Path:
        """Экспортирует результат в DOCX"""
        try:
            if output_path is None:
                outputs_dir = self.project_root / "outputs"
                outputs_dir.mkdir(exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = outputs_dir / f"consultation_{timestamp}.docx"
            
            return self.exporter.create_report(result, output_path)
        except Exception as e:
            print(f"❌ Ошибка экспорта: {e}")
            return None


class DocxExporter:
    """Экспортер результатов в DOCX формат"""
    
    def __init__(self, project_root: Path):
        """Инициализация экспортера"""
        self.project_root = project_root
        self.outputs_dir = project_root / "outputs"
        self.outputs_dir.mkdir(exist_ok=True)
    
    def create_report(self, result: dict, output_path: Path) -> Path:
        """Создает отчет в DOCX формате"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('Юридическая консультация', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата и время
        timestamp = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        doc.add_paragraph(f"Дата: {timestamp}")
        
        # Вопрос
        doc.add_heading('Вопрос', level=1)
        question = result.get('question', 'Не указан')
        doc.add_paragraph(question)
        
        # Категория
        doc.add_heading('Категория', level=1)
        category = result.get('category', 'Не определена')
        doc.add_paragraph(category)
        
        # Ответ
        doc.add_heading('Ответ', level=1)
        answer = result.get('answer_generation', {}).get('answer', 'Ответ не получен')
        doc.add_paragraph(answer)
        
        # Источники
        doc.add_heading('Использованные источники', level=1)
        articles = result.get('found_articles', [])
        if articles:
            for article in articles:
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                article_text = f"• Статья {article.get('article', 'N/A')} - {law_display}"
                doc.add_paragraph(article_text)
        else:
            doc.add_paragraph("Источники не найдены")
        
        # Сохраняем документ
        doc.save(str(output_path))
        print(f"✅ Отчет сохранен: {output_path}")
        
        return output_path


def run_telegram_bot():
    """Запускает Telegram-бота"""
    try:
        from telegram import Update
        from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
    except ImportError:
        print("❌ Требуется установить python-telegram-bot: pip install python-telegram-bot")
        return
    
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    if not token:
        print("❌ Переменная окружения TELEGRAM_BOT_TOKEN не установлена")
        return
    
    system = LegalAISystem(auto_init_rag=True)
    
    async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик команды /start"""
        await update.message.reply_text(
            "👋 Добро пожаловать в ИИ Юриста!\n\n"
            "Отправьте мне ваш юридический вопрос, и я помогу вам найти ответ в законодательстве."
        )
    
    async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик текстовых сообщений"""
        text = update.message.text
        
        await update.message.reply_text("⏳ Анализирую ваш вопрос...")
        
        try:
            result = await system.process_question(text)
            
            if result.get("status") == "error":
                await update.message.reply_text(f"❌ Ошибка: {result.get('error', 'Неизвестная ошибка')}")
                return
            
            answer = result.get('answer_generation', {}).get('answer', 'Ответ не получен')
            articles = result.get('found_articles', [])
            
            response = f"📋 Ответ:\n\n{answer}\n\n"
            
            if articles:
                response += "📚 Источники:\n"
                for article in articles[:5]:
                    law_filename = article.get('law', 'N/A')
                    law_display = get_law_display_name(law_filename)
                    response += f"• Статья {article.get('article', 'N/A')} - {law_display}\n"
            
            await update.message.reply_text(response)
            
        except Exception as e:
            await update.message.reply_text(f"❌ Ошибка обработки: {str(e)}")
    
    # Создаем приложение
    app = Application.builder().token(token).build()
    
    # Добавляем обработчики
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    
    # Запускаем бота
    print("🤖 Telegram-бот запущен. Нажмите Ctrl+C для остановки.")
    app.run_polling()


def main():
    """Главная функция"""
    parser = argparse.ArgumentParser(description="ИИ Юрист - система юридических консультаций")
    parser.add_argument("--index", action="store_true", help="Запустить AI-индексацию законов")
    parser.add_argument("--telegram", action="store_true", help="Запустить в режиме Telegram-бота")
    parser.add_argument("--question", type=str, help="Задать вопрос напрямую")
    
    args = parser.parse_args()
    
    # Режим индексации
    if args.index:
        system = LegalAISystem(auto_init_rag=False)
        asyncio.run(system.index_laws_with_ai())
        return
    
    # Режим Telegram-бота
    if args.telegram:
        run_telegram_bot()
        return
    
    # Режим прямого вопроса
    if args.question:
        system = LegalAISystem()
        result = asyncio.run(system.process_question(args.question))
        
        if result.get("status") == "error":
            print(f"❌ Ошибка: {result.get('error')}")
            return
        
        print("\n" + "="*60)
        print("📋 ОТВЕТ:")
        print("="*60)
        answer = result.get('answer_generation', {}).get('answer', 'Ответ не получен')
        print(answer)
        
        articles = result.get('found_articles', [])
        if articles:
            print("\n📚 ИСТОЧНИКИ:")
            for article in articles:
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
        if articles:
            print("\n📚 ИСТОЧНИКИ:")
            for article in articles:
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
        if articles:
            print("\n📚 ИСТОЧНИКИ:")
            for article in articles:
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
        if articles:
            print("\n📚 ИСТОЧНИКИ:")
            for article in articles:
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
        if articles:
            print("\n📚 ИСТОЧНИКИ:")
            for article in articles:
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
                law_filename = article.get('law', 'N/A')
                law_display = get_law_display_name(law_filename)
                print(f"  • Статья {article.get('article', 'N/A')} - {law_display}")
        
        export = input("\nЭкспортировать в DOCX? (y/n): ").lower()
        if export == 'y':
            path = system.export_to_docx(result)
            if path:
                print(f"✅ Файл сохранен: {path}")
        return
    
    # Интерактивный режим
    print("="*60)
    print("🏛️  ИИ ЮРИСТ - Интерактивный режим")
    print("="*60)
    print("Введите 'выход' или 'exit' для завершения\n")
    
    system = LegalAISystem()
    
    while True:
        try:
            question = input("\n❓ Ваш вопрос: ").strip()
            
            if question.lower() in ['выход', 'exit', 'quit']:
                print("👋 До свидания!")
                break
            
            if not question:
                continue
            
            result = asyncio.run(system.process_question(question))
            
            if result.get("status") == "error":
                print(f"❌ Ошибка: {result.get('error')}")
                continue
            
            print("\n" + "="*60)
            print("📋 ОТВЕТ:")
            print("="*60)
            answer = result.get('answer_generation', {}).get('answer', 'Ответ не получен')
            print(answer)
            
            articles = result.get('found_articles', [])
            if articles:
                print("\n📚 ИСТОЧНИКИ:")
                for article in articles:
                    print(f"  • Статья {article.get('article', 'N/A')} - {article.get('law', 'N/A')}")
            
            export = input("\nЭкспортировать в DOCX? (y/n): ").lower()
            if export == 'y':
                path = system.export_to_docx(result)
                if path:
                    print(f"✅ Файл сохранен: {path}")
        
        except KeyboardInterrupt:
            print("\n👋 До свидания!")
            break
        except Exception as e:
            print(f"❌ Ошибка: {e}")


if __name__ == "__main__":
    main()