# src/core/sequential_searcher.py
import json
import re
import time
import asyncio
import os           # 🔹 НОВОЕ
import psutil       # 🔹 НОВОЕ
from src.core.simple_law_parser import SimpleLawParser # импорт парсера(для симантики)
from pathlib import Path
from typing import List, Dict
from difflib import SequenceMatcher
from rapidfuzz import fuzz, process
from api_manager import api_manager # импорт файла где все api
from src.core.answer_generator import AnswerGenerator # импорт генератора ответа
from src.core.docx_intelligent_parser import DocxIntelligentParser # импорт парсера для чанкирования законов для ИИ

class RAG2Validator:
    """RAG2: Точная валидация статей, найденных ИИ-поисковиком"""
    
    def __init__(self, laws_dir: Path):
        self.laws_dir = laws_dir
        self.law_cache = {}
        self.validation_log = []
        
    def validate_articles(self, articles: List[Dict], original_question: str = "") -> Dict:
        """Валидирует статьи, найденные ИИ-поисковиком"""
        print(f"🎯 RAG2: Запуск точной валидации {len(articles)} статей")
        self.validation_log.append({
            "timestamp": time.time(),
            "action": "start_validation",
            "articles_count": len(articles),
            "question": original_question
        })
        
        validation_results = {
            "total_articles_received": len(articles),
            "articles_validated": 0,
            "articles_corrected": 0,
            "articles_rejected": 0,
            "validation_details": [],
            "final_articles": [],
            "validation_log": self.validation_log
        }
        
        for i, article in enumerate(articles):
            print(f"   📋 RAG2: Валидация статьи {i+1}/{len(articles)}")
            print(f"      📖 Статья: {article.get('article', 'Unknown')}")
            print(f"      📚 Закон: {article.get('law', 'Unknown')}")
            
            validation_result = self._validate_single_article(article, original_question)
            validation_results["validation_details"].append(validation_result)
            
            if validation_result["status"] == "validated_exact":
                validation_results["articles_validated"] += 1
                validation_results["final_articles"].append(validation_result["validated_article"])
                print(f"      ✅ Точное соответствие")
            elif validation_result["status"] == "validated_corrected":
                validation_results["articles_corrected"] += 1
                validation_results["final_articles"].append(validation_result["validated_article"])
                print(f"      🔧 Исправлено: {validation_result.get('correction_note', '')}")
            else:
                validation_results["articles_rejected"] += 1
                print(f"      ❌ Отклонена: {validation_result['reason']}")
        
        print(f"✅ RAG2: Валидация завершена")
        print(f"   ✅ Валидировано: {validation_results['articles_validated']}")
        print(f"   🔧 Исправлено: {validation_results['articles_corrected']}") 
        print(f"   ❌ Отклонено: {validation_results['articles_rejected']}")
        
        self.validation_log.append({
            "timestamp": time.time(),
            "action": "end_validation",
            "results": {
                "validated": validation_results["articles_validated"],
                "corrected": validation_results["articles_corrected"],
                "rejected": validation_results["articles_rejected"]
            }
        })
        
        return validation_results
    
    def _validate_single_article(self, article: Dict, question: str) -> Dict:
        """Валидирует одну статью с детальным логированием"""
        law_name = article.get('law', '')
        article_title = article.get('article', '')
        article_content = article.get('content', '')
        
        # 🔥 ДОБАВЛЯЕМ: Расширенное логирование для отладки
        print(f"      🔍 RAG2: Начало валидации статьи:")
        print(f"          Название: {article_title}")
        print(f"          Закон: {law_name}")
        print(f"          Длина контента: {len(article_content) if article_content else 0} символов")
    
        # 🔥 КРИТИЧЕСКОЕ ЛОГИРОВАНИЕ: Проверка наличия названия закона
        if not law_name:
            print(f"      🚨 КРИТИЧЕСКОЕ ПРЕДУПРЕЖДЕНИЕ: Статья '{article_title}' передана в валидатор БЕЗ названия закона!")
            print(f"      Проверьте цепочку: Глава -> Анализ ИИ -> Передача в валидатор.")
            print(f"      Полная статья для отладки: {json.dumps(article, ensure_ascii=False, indent=2)}")
            
            # Записываем в debug-файл для дальнейшего анализа
            debug_file = self.laws_dir.parent / "debug_missing_law.txt"
            try:
                with open(debug_file, 'a', encoding='utf-8') as f:
                    f.write(f"\n{'='*80}\n")
                    f.write(f"Время: {time.ctime()}\n")
                    f.write(f"Статья без закона: {article_title}\n")
                    f.write(f"Полный объект статьи:\n")
                    f.write(json.dumps(article, ensure_ascii=False, indent=2))
                    f.write(f"\n{'='*80}\n")
                print(f"      🔍 Информация сохранена в: {debug_file}")
            except Exception as debug_e:
                print(f"      ⚠️ Не удалось сохранить debug-информацию: {debug_e}")        
        
        
        validation_step = {
            "article_title": article_title,
            "law_name": law_name,
            "steps": []
        }
        
        # Список файлов, которые мы считаем валидными "как есть", без проверки на Статьи/Пункты.
        WHITELIST_FILES = {
            "postanov o neusto k dolevoi.docx",
            "Plenum on the gk 2.docx", 
            # Добавляйте свои файлы сюда
        }
        
        # Если файл в белом списке и у него есть контент -> СРАЗУ ВАЛИДИРУЕМ
        if law_name in WHITELIST_FILES and article_content:
             print(f"      ✅ RAG2: Файл '{law_name}' в белом списке. Пропуск строгой валидации.")
             return {
                "status": "validated_exact", # Или validated_corrected, не принципиально
                "reason": "Файл в белом списке (валидация отключена)",
                "original_article": article,
                "validated_article": article, # Возвращаем как есть
                "match_confidence": 1.0,
                "validation_steps": validation_step
            }        
        
        
        
        # Шаг 1: Базовые проверки
        validation_step["steps"].append({
            "step": "basic_checks",
            "input": {"law_name": law_name, "has_content": bool(article_content)}
        })
        
        # 🔥 УЛУЧШАЕМ: Более детальное сообщение об ошибке
        if not law_name or not article_content:
            reason = "Отсутствует название закона или содержание статьи"
            if not law_name:
                reason = f"Отсутствует название закона. Статья: '{article_title}'"
                print(f"      ❌ RAG2: {reason}")
            elif not article_content:
                reason = f"Отсутствует содержание статьи. Закон: '{law_name}', Статья: '{article_title}'"
                print(f"      ❌ RAG2: {reason}")
            
            validation_step["steps"][-1]["result"] = "fail"
            validation_step["steps"][-1]["reason"] = reason
            return {
                "status": "rejected",
                "reason": reason,
                "original_article": article,
                "validation_steps": validation_step
            }
        
        validation_step["steps"][-1]["result"] = "pass"        
        
        # Шаг 2: Загрузка закона
        validation_step["steps"].append({
            "step": "load_law",
            "input": {"law_name": law_name}
        })
        
        law_content = self._load_law_content(law_name)
        if not law_content:
            validation_step["steps"][-1]["result"] = "fail"
            validation_step["steps"][-1]["reason"] = f"Закон '{law_name}' не найден"
            return {
                "status": "rejected", 
                "reason": f"Закон '{law_name}' не найден",
                "original_article": article,
                "validation_steps": validation_step
            }
        
        validation_step["steps"][-1]["result"] = "pass"
        validation_step["steps"][-1]["law_size"] = len(law_content)
        
        # Шаг 3: Поиск точного соответствия
        validation_step["steps"].append({
            "step": "exact_match_search",
            "input": {"content_length": len(article_content)}
        })
        
        exact_match = self._find_exact_article_match(article_content, law_content, law_name)
        if exact_match:
            validation_step["steps"][-1]["result"] = "success"
            validation_step["steps"][-1]["match_type"] = "exact"
            return {
                "status": "validated_exact",
                "reason": "Найдено точное соответствие в законе",
                "original_article": article,
                "validated_article": exact_match,
                "match_confidence": 1.0,
                "validation_steps": validation_step
            }
        
        validation_step["steps"][-1]["result"] = "fail"
        
        # Шаг 4: Поиск с коррекцией
        validation_step["steps"].append({
            "step": "corrected_match_search",
            "input": {"article_title": article_title}
        })
        
        corrected_match = self._find_corrected_article_match(article_title, article_content, law_content, law_name)
        if corrected_match:
            validation_step["steps"][-1]["result"] = "success"
            validation_step["steps"][-1]["match_type"] = "corrected"
            validation_step["steps"][-1]["confidence"] = corrected_match.get("match_confidence", 0.8)
            return {
                "status": "validated_corrected",
                "reason": "Статья найдена с исправлением неточностей",
                "original_article": article,
                "validated_article": corrected_match,
                "match_confidence": corrected_match.get("match_confidence", 0.8),
                "validation_steps": validation_step
            }
        
        validation_step["steps"][-1]["result"] = "fail"
        
        # Шаг 5: Проверка на выдуманные статьи
        validation_step["steps"].append({
            "step": "fabrication_check",
            "input": {"question": question}
        })
        
        if self._is_likely_fabricated(article, law_content, question):
            validation_step["steps"][-1]["result"] = "fail"
            validation_step["steps"][-1]["reason"] = "Статья вероятно выдумана ИИ"
            return {
                "status": "rejected",
                "reason": "Статья вероятно выдумана ИИ",
                "original_article": article,
                "validation_steps": validation_step
            }
        
        validation_step["steps"][-1]["result"] = "pass"
        
        return {
            "status": "rejected",
            "reason": "Не удалось найти соответствие в законе",
            "original_article": article,
            "validation_steps": validation_step
        }
    
    def _load_law_content(self, law_name: str) -> str:
        """Загружает содержание закона из файла (поддерживает .txt и .docx)"""
        if law_name in self.law_cache:
            return self.law_cache[law_name]
        
        law_path = self.laws_dir / law_name
        if not law_path.exists():
            print(f"      ⚠️ RAG2: Файл закона не найден: {law_path}")
            return None
        
        try:
            content = ""
            # Проверка расширения файла
            if law_name.lower().endswith('.docx'):
                # Чтение DOCX
                import docx
                doc = docx.Document(law_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                content = "\n".join(full_text)
                print(f"      ✅ RAG2: DOCX закон загружен: {law_name} ({len(content)} символов)")
            else:
                # Чтение TXT (с попыткой разных кодировок)
                try:
                    with open(law_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    with open(law_path, 'r', encoding='windows-1251') as f:
                        content = f.read()
                print(f"      ✅ RAG2: TXT закон загружен: {law_name} ({len(content)} символов)")

            if content:
                self.law_cache[law_name] = content
                return content
            return None

        except Exception as e:
            print(f"      ❌ RAG2: Ошибка чтения закона {law_name}: {e}")
            return None
    
    def _find_exact_article_match(self, article_content: str, law_content: str, law_name: str) -> Dict:
        """Ищет точное соответствие статьи в законе"""
        normalized_article = self._normalize_text(article_content)
        
        if normalized_article in law_content:
            start_idx = law_content.find(normalized_article)
            end_idx = start_idx + len(normalized_article)
            article_title = self._extract_article_title_from_context(law_content, start_idx)
            
            print(f"      ✅ RAG2: Найдено точное соответствие для '{article_title}'")
            return {
                "law": law_name,
                "article": article_title,
                "content": article_content,
                "exact_match": True,
                "validation_method": "exact_match"
            }
        
        print(f"      ❌ RAG2: Точное соответствие не найдено")
        return None
    
    def _find_corrected_article_match(self, article_title: str, article_content: str, 
                                    law_content: str, law_name: str) -> Dict:
        """Ищет статью с коррекцией небольших неточностей"""
        
        # Стратегия 1: Поиск по заголовку статьи
        print(f"      🔍 RAG2: Поиск по заголовку: {article_title}")
        title_match = self._find_article_by_title(article_title, law_content, law_name)
        if title_match:
            return title_match
        
        # Стратегия 2: Fuzzy matching по содержанию
        print(f"      🔍 RAG2: Fuzzy поиск по содержанию")
        content_match = self._find_article_by_fuzzy_content(article_content, law_content, law_name)
        if content_match:
            return content_match
        
        # Стратегия 3: Поиск по ключевым фразам
        print(f"      🔍 RAG2: Поиск по ключевым фразам")
        phrase_match = self._find_article_by_key_phrases(article_content, law_content, law_name)
        if phrase_match:
            return phrase_match
        
        print(f"      ❌ RAG2: Корректирующее соответствие не найдено")
        return None
    
    def _find_article_by_title(self, article_title: str, law_content: str, law_name: str) -> Dict:
        """Ищет статью по заголовку с коррекцией номеров"""
        article_number = self._extract_article_number(article_title)
        if not article_number:
            return None
        
        print(f"      🔍 RAG2: Поиск статьи с номером: {article_number}")
        
        article_patterns = [
            rf'Статья\s+{article_number}[\.\s]',
            rf'СТАТЬЯ\s+{article_number}[\.\s]',
            rf'ст\.\s*{article_number}',
            rf'Статья\s+{self._normalize_article_number(article_number)}[\.\s]'
        ]
        
        for pattern in article_patterns:
            match = re.search(pattern, law_content, re.IGNORECASE)
            if match:
                article_text = self._extract_article_by_position(law_content, match.start())
                if article_text:
                    print(f"      ✅ RAG2: Найдена статья по номеру: {article_number}")
                    return {
                        "law": law_name,
                        "article": f"Статья {article_number}",
                        "content": article_text,
                        "exact_match": False,
                        "validation_method": "title_correction",
                        "match_confidence": 0.9,
                        "correction_note": f"Скорректирован номер статьи: {article_number}"
                    }
        
        return None
    
    def _find_article_by_fuzzy_content(self, article_content: str, law_content: str, law_name: str) -> Dict:
        """Ищет статью по fuzzy matching содержания"""
        articles = self._split_law_into_articles(law_content, law_name)
        if not articles:
            return None
        
        best_match = None
        best_score = 0
        
        for law_article in articles:
            score = fuzz.ratio(
                self._normalize_text(article_content[:500]),
                self._normalize_text(law_article['content'][:500])
            )
            
            if score > best_score and score > 70:
                best_score = score
                best_match = law_article
        
        if best_match:
            print(f"      ✅ RAG2: Найдена статья по схожести ({best_score}%): {best_match['title']}")
            return {
                "law": law_name,
                "article": best_match['title'],
                "content": best_match['content'],
                "exact_match": False,
                "validation_method": "fuzzy_content",
                "match_confidence": best_score / 100,
                "correction_note": f"Найдено по схожести содержания ({best_score}%)"
            }
        
        return None
    
    def _find_article_by_key_phrases(self, article_content: str, law_content: str, law_name: str) -> Dict:
        """Ищет статью по ключевым юридическим фразам"""
        key_phrases = self._extract_key_legal_phrases(article_content)
        if not key_phrases:
            return None
        
        print(f"      🔍 RAG2: Ключевые фразы: {key_phrases}")
        
        articles = self._split_law_into_articles(law_content, law_name)
        best_article = None
        best_phrase_count = 0
        
        for law_article in articles:
            phrase_count = sum(1 for phrase in key_phrases if phrase in law_article['content'])
            if phrase_count > best_phrase_count:
                best_phrase_count = phrase_count
                best_article = law_article
        
        if best_article and best_phrase_count >= 2:
            print(f"      ✅ RAG2: Найдена статья по {best_phrase_count} ключевым фразам: {best_article['title']}")
            return {
                "law": law_name,
                "article": best_article['title'],
                "content": best_article['content'],
                "exact_match": False,
                "validation_method": "key_phrases",
                "match_confidence": min(0.8, best_phrase_count / len(key_phrases)),
                "correction_note": f"Найдено по {best_phrase_count} ключевым фразам"
            }
        
        return None
    
    def _is_likely_fabricated(self, article: Dict, law_content: str, question: str) -> bool:
        """Проверяет, является ли статья вероятно выдуманной (без проверки релевантности вопросу)"""
        content = article.get('content', '').lower()
        
        # 1. Проверка на "общие фразы" (оставляем, это полезно против воды)
        fabrication_indicators = [
            "в соответствии с законодательством",
            "установлено законом", 
            "предусмотрено нормативными актами",
            "срок определяется соглашением сторон",
            "размер устанавливается в договоре",
            "не менее и не более", 
            "одновременно и то и другое",
            "очень важно", "надо помнить", "следует учитывать"
        ]
        
        general_phrase_count = sum(1 for phrase in fabrication_indicators if phrase in content)
        if general_phrase_count >= 2:
            print(f"      ⚠️ RAG2: Обнаружено {general_phrase_count} общих фраз - возможная выдумка")
            return True
        
        # 2. Проверка на структуру (оставляем, полезно против обычного текста вместо статей)
        if not self._has_legal_article_structure(content):
            print(f"      ⚠️ RAG2: Отсутствует структура юридической статьи")
            return True
        
        # 3. Проверка релевантности вопросу (УДАЛЯЕМ/ОТКЛЮЧАЕМ)
        # if question and not self._is_relevant_to_question(content, question):
        #    print(f"      ⚠️ RAG2: Статья не релевантна вопросу")
        #    return True
        
        return False
    
    def _split_law_into_articles(self, law_content: str, law_name: str) -> List[Dict]:
        """Разбивает закон на отдельные статьи"""
        articles = []
        
        article_patterns = [
            r'(Статья\s+\d+[\.\d]*[^\n]*)\n',
            r'(СТАТЬЯ\s+\d+[\.\d]*[^\n]*)\n',
            r'(ст\.\s*\d+[\.\d]*[^\n]*)\n'
            r'(Пункт\s+\d+[\.\d]*[^\n]*)\n',   # <-- NEW
            r'(^\d+\.\s+[^\n]*)\n'             # <-- NEW: "1. Установить..." (в начале строки)
        ]
        
        for pattern in article_patterns:
            matches = list(re.finditer(pattern, law_content, re.IGNORECASE))
            for i, match in enumerate(matches):
                title = match.group(1).strip()
                start_pos = match.end()
                end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(law_content)
                content = law_content[start_pos:end_pos].strip()
                
                if len(content) > 5000:
                    content = content[:5000] + "..."
                
                articles.append({
                    'title': title,
                    'content': content,
                    'law': law_name
                })
        
        return articles
    
    def _normalize_text(self, text: str) -> str:
        """Нормализует текст для сравнения"""
        if not text:
            return ""
        text = text.lower().strip()
        text = re.sub(r'[^\w\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text
    
    def _extract_article_number(self, article_title: str) -> str:
        """Извлекает номер статьи из заголовок"""
        patterns = [
            r'Статья\s+(\d+[\.\d]*)',
            r'СТАТЬЯ\s+(\d+[\.\d]*)', 
            r'ст\.\s*(\d+[\.\d]*)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, article_title, re.IGNORECASE)
            if match:
                return match.group(1)
        
        return None
    
    def _normalize_article_number(self, number: str) -> str:
        """Нормализует номер статьи"""
        return re.sub(r'\.$', '', number)
    
    def _extract_article_title_from_context(self, law_content: str, position: int) -> str:
        """Извлекает заголовок статьи из контекста"""
        start = max(0, position - 200)
        context = law_content[start:position]
        
        patterns = [
            r'(Статья\s+\d+[\.\d]*[^\n]*)',
            r'(СТАТЬЯ\s+\d+[\.\d]*[^\n]*)'
        ]
        
        for pattern in patterns:
            matches = list(re.finditer(pattern, context, re.IGNORECASE))
            if matches:
                return matches[-1].group(1).strip()
        
        return "Статья (заголовок не определен)"
    
    def _extract_article_by_position(self, law_content: str, start_position: int) -> str:
        """Извлекает статью по начальной позиции"""
        next_article_pattern = r'\n(Статья\s+\d+|СТАТЬЯ\s+\d+)'
        next_match = re.search(next_article_pattern, law_content[start_position + 10:], re.IGNORECASE)
        
        if next_match:
            end_position = start_position + 10 + next_match.start()
            return law_content[start_position:end_position].strip()
        else:
            return law_content[start_position:].strip()
    
    def _extract_key_legal_phrases(self, content: str) -> List[str]:
        """Извлекает ключевые юридические фразы"""
        legal_phrases = [
            "вправе требовать", "обязан выплатить", "срок действия", 
            "расторжение договора", "ответственность сторон", "порядок осуществления",
            "в установленном порядке", "в соответствии с", "имеет право", 
            "обязан предоставить", "не вправе", "запрещается", "разрешается"
        ]
        
        found_phrases = []
        content_lower = content.lower()
        
        for phrase in legal_phrases:
            if phrase in content_lower:
                found_phrases.append(phrase)
        
        return found_phrases
    
    def _has_legal_article_structure(self, content: str) -> bool:
        """Проверяет, имеет ли текст структуру юридической статьи"""
        structure_indicators = [
            r'\d+\.', r'[а-я]\)', r'—\s', 'часть', 'пункт', 'подпункт'
        ]
        
        indicators_count = sum(1 for pattern in structure_indicators if re.search(pattern, content.lower()))
        return indicators_count >= 2
    
    def _is_relevant_to_question(self, content: str, question: str) -> bool:
        """Проверяет релевантность статьи вопросу"""
        question_words = set(re.findall(r'\b[а-я]{4,}\b', question.lower()))
        content_words = set(re.findall(r'\b[а-я]{4,}\b', content.lower()))
        
        common_words = question_words & content_words
        return len(common_words) >= 2

class IntelligentChunkManager:
    """Управляет разбивкой больших глав на подчанки для избежания ошибки 413"""

    def __init__(self, max_chunk_size: int = 120000):
        self.max_chunk_size = max_chunk_size

    def prepare_chapter_analysis(self, chapter: Dict) -> List[Dict]:
        """
        Подготавливает блок (чанк) для анализа.
        Так как парсер (v3) уже нарезал текст на умные чанки, 
        здесь мы просто возвращаем его как 'готовый к работе'.
        """
        content = chapter.get('content', '')
        title = chapter.get('chapter_title', '') or chapter.get('title', '')
        
        # Просто логируем для отладки
        # print(f" 📦 Обработка готового чанка '{title}' ({len(content)} симв.)")
        
        return [{
            'type': chapter.get('type', 'full_node'), # Берем тип от парсера или дефолт
            'title': title,
            'content': content,
            'chunk_index': 0, # У парсера v3 чанки независимы, индексы не критичны тут
            'total_chunks': 1,
            'law': chapter.get('law', ''),
            'original_chapter': title,
            'level': chapter.get('level', 0)
        }]

    def _split_chapter_into_chunks(self, chapter: Dict, content: str) -> List[Dict]:
        """
        Разбивает большую главу на смысловые подчанки (Параграфы, Разделы).
        Если структурных блоков нет, режет по границам статей.
        """
        chunks_content = []
        
        # 1. Попытка найти крупные структуры второго уровня (Параграфы, Разделы)
        # Приоритет: Параграф -> Раздел -> Статья
        structure_patterns = [
            r'\n(?:Параграф|§)\s+\d+',        # § 1. ...
            r'\n(?:Раздел|РАЗДЕЛ)\s+[IVX]+',  # Раздел I ...
            r'\n(?:Статья|СТАТЬЯ)\s+\d+'      # Статья 1 ... (как fallback)
        ]
        
        boundaries = []
        found_pattern = None
        
        for pattern in structure_patterns:
            found = list(re.finditer(pattern, content))
            # Если нашли хотя бы 2 блока (чтобы было что делить) или 1 блок в середине
            if len(found) >= 1:
                print(f"  🎯 Найдена внутренняя структура по паттерну: {pattern}")
                boundaries = [m.start() for m in found]
                found_pattern = pattern
                break
        
        # Если границ не нашли вообще - режем просто по размеру (совсем крайний случай)
        if not boundaries:
            print("  ⚠️ Структура не найдена, режу по размеру.")
            chunks_content = self._split_by_size(content)
        else:
            # Формируем чанки по найденным границам
            # Добавляем начало текста (до первого параграфа) если оно есть
            if boundaries[0] > 0:
                chunks_content.append(content[:boundaries[0]])
                
            for i in range(len(boundaries)):
                start = boundaries[i]
                end = boundaries[i+1] if i+1 < len(boundaries) else len(content)
                
                block_content = content[start:end]
                
                # Рекурсивная проверка: если даже этот блок (Параграф) больше лимита
                if len(block_content) > self.max_chunk_size:
                    print(f"  ⚠️ Блок {i+1} слишком большой ({len(block_content)}), режу его дополнительно...")
                    # Тут можно было бы искать Статьи, но для простоты режем по размеру/абзацам
                    sub_chunks = self._split_by_size(block_content)
                    chunks_content.extend(sub_chunks)
                else:
                    chunks_content.append(block_content)

        # 2. Упаковываем тексты в словари с метаданными
        prepared_chunks = []
        total_chunks = len(chunks_content)
        
        for i, chunk_text in enumerate(chunks_content):
            # Добавляем "Контекстный Якорь" в начало текста, чтобы ИИ не потерялся
            # Ищем заголовок внутри чанка (первая строка)
            first_line = chunk_text.strip().split('\n')[0][:100]
            chunk_title_display = f"{chapter['title']} -> {first_line}..."
            
            prepared_chunks.append({
                'type': 'chapter_chunk',
                'title': f"{chapter['title']} [Чанк {i+1}/{total_chunks}]",
                'content': chunk_text, # Текст блока как есть
                'chunk_index': i,
                'total_chunks': total_chunks,
                'original_chapter': chapter['title'],
                'law': chapter.get('law', '')
            })

        print(f" 📦 Итог: {total_chunks} смысловых чанков.")
        return prepared_chunks

    def _split_by_size(self, content: str) -> List[str]:
        """Разбивает контент просто по размеру (fallback), стараясь не резать предложения"""
        chunks = []
        current_pos = 0
        total_len = len(content)
        
        while current_pos < total_len:
            target_end = current_pos + self.max_chunk_size
            
            if target_end >= total_len:
                chunks.append(content[current_pos:])
                break
            
            # Ищем ближайший конец абзаца или предложения перед лимитом
            # Смотрим окно в 500 символов назад от лимита
            search_window = content[max(current_pos, target_end - 2000) : target_end]
            
            # Приоритеты разделителей
            last_break = -1
            for sep in ['\n\n', '\n', '. ']:
                last_break = search_window.rfind(sep)
                if last_break != -1:
                    # Корректируем позицию относительно начала search_window
                    abs_break = max(current_pos, target_end - 2000) + last_break + len(sep)
                    chunks.append(content[current_pos:abs_break])
                    current_pos = abs_break
                    break
            
            if last_break == -1:
                # Если совсем ничего не нашли (огромный кусок текста без точек), режем жестко
                chunks.append(content[current_pos:target_end])
                current_pos = target_end
                
        return chunks

class SequentialArticleSearcher:
    """Автономный поисковик статей, который самостоятельно находит и анализирует законы"""
    
    def __init__(self, navigator=None, project_root=None):
        self.navigator = navigator
        # Определяем корневую директорию проекта
        if project_root:
            self.project_root = Path(project_root)
        else:
            # Поднимаемся на 2 уровня вверх из src/core/
            self.project_root = Path(__file__).parent.parent.parent
        
        self._process = psutil.Process(os.getpid())
        
        self.laws_dir = self.project_root / "laws"
        print(f"   📁 Корень проекта: {self.project_root}")
        print(f"   📁 Папка законов: {self.laws_dir}")
        print(f"   📁 Существует: {self.laws_dir.exists()}")
        
        self.simple_parser = SimpleLawParser()
        
        # Инициализируем RAG2 валидатор
        self.rag2_validator = RAG2Validator(self.laws_dir)
        
        self.chunk_manager = IntelligentChunkManager(max_chunk_size=120000)
        
        self.search_hints = self._load_search_hints()        
        
        self.search_log = {
            "question": "",
            "category": "",
            "total_chapters_analyzed": 0,
            "chapters_processed": [],
            "articles_found": [],
            "search_sequence": [],
            "start_time": 0,
            "end_time": 0,
            "sufficiency_assessment": "",
            "expansion_triggered": False,
            "laws_searched": [],
            "rag2_validation": {}  # Добавляем поле для результатов RAG2
        }
        self.articles_file_path = None
        self.processed_chapters = set()
    
    def _load_search_hints(self) -> Dict:
        """Загружает умные советы из JSON файла"""
        try:
            # Путь к файлу (настройте под свою структуру)
            hints_path = Path(__file__).parent.parent.parent / "search_hints.json" 
            
            if hints_path.exists():
                with open(hints_path, 'r', encoding='utf-8') as f:
                    hints = json.load(f)
                print(f" ✅ Умные советы загружены из {hints_path}")
                return hints
            else:
                print(f" ⚠️ Файл search_hints.json не найден по пути {hints_path}. Использую пустые советы.")
                return {}
        except Exception as e:
            print(f" ❌ Ошибка загрузки search_hints.json: {e}")
            return {}    
    
    async def sequential_search(self, question: str, category: str,
                            output_dir: str = "search_results") -> Dict:
        """Автономный поиск: сам находит законы, главы и статьи с RAG2 валидацией"""

        print(f"🚀 АВТОНОМНЫЙ ПОИСК: '{question}'")
        print(f"   📂 Категория: {category}")

        # 🔹 Старт: время и ресурсы на ВЕСЬ процесс поиска
        search_start_time = time.time()
        resource_stats = {}
        try:
            proc = psutil.Process(os.getpid())
            cpu_start = proc.cpu_times()
            mem_start = proc.memory_info().rss  # байты
        except Exception:
            proc = None
            cpu_start = None
            mem_start = None

        # Создаем папку для результатов в корне проекта
        results_dir = self.project_root / output_dir
        results_dir.mkdir(exist_ok=True)

        # Создаем уникальные файлы для этого поиска
        timestamp = int(time.time())

        # --- ИЗМЕНЕНИЕ: меняем расширение на .json ---
        self.articles_file_path = results_dir / f"articles_{timestamp}.json"
        log_file_path = results_dir / f"search_log_{timestamp}.json"

        self.search_log = {
            "question": question,
            "category": category,
            "total_chapters_analyzed": 0,
            "chapters_processed": [],
            "articles_found": [],
            "search_sequence": [],
            "start_time": time.time(),
            "end_time": 0,
            "sufficiency_assessment": "",
            "expansion_triggered": False,
            "articles_file": str(self.articles_file_path),
            "log_file": str(log_file_path),
            "laws_searched": [],
            "rag2_validation": {}  # Добавляем поле для результатов RAG2
        }

        # --- ИЗМЕНЕНИЕ: Блок инициализации текстового файла больше не нужен ---
        # Мы создадим JSON файл целиком перед этапом генерации ответа
        # with open(self.articles_file_path, "w", encoding="utf-8") as f:
        #     f.write(f"# Найденные статьи для вопроса: {question}\n")
        #     f.write(f"# Категория: {category}\n")
        #     f.write(f"# Время начала: {time.ctime()}\n")
        #     f.write("=" * 80 + "\n\n")

        print(f"   📁 Файл для статей: {self.articles_file_path}")

        # ... (дальше идет try/except блок без изменений,
        # но НЕ ЗАБУДЬТЕ добавить вызов _save_final_articles_to_json перед генерацией,
        # как мы обсуждали ранее)

        try:
            # Этап 1: Получаем законы для категории
            law_files = self._get_laws_for_category(category)
            if not law_files:
                return self._create_error_result(
                    "Не найдены законы для указанной категории",
                    log_file_path,
                )

            print(f"   📚 Законы для поиска: {len(law_files)} файлов")

            # Этап 2: Находим релевантные главы через навигатор
            # Этап 2: Находим релевантные контейнеры (глав/разделов) через RAG-1... (НОВАЯ ЛОГИКА)
            print("   🔍 Этап 2: Поиск релевантных контейнеров (глав/разделов) через RAG-1...")

            if self.navigator:
                # Вызываем новый метод. Он вернет список словарей с 'content'
                chapters_data = self.navigator.find_relevant_containers(
                    question,
                    law_files,
                    top_k=10,
                )

                chapters = []
                for container in chapters_data:
                    # Просто перекладываем данные, сохраняя content
                    chapters.append({
                        "title": container['title'],
                        "chapter_title": container['title'],
                        "law": container['law'],
                        "content": container['content'],  # <-- ВАЖНО: Полный текст главы
                        "type": container['type'],
                    })
            else:
                print("   ⚠️ Навигатор не подключен.")
                chapters = []

            if not chapters:
                return self._create_error_result(
                    "Не найдено релевантных глав (контейнеров)",
                    log_file_path,
                )

            print(f"   🔍 Найдено контейнеров для анализа: {len(chapters)}")

            # Этап 3: Анализ глав и поиск статей (initial, только главы от RAG1)
            articles_initial_raw = await self._analyze_chapters_phase(
                question,
                chapters,
                category,
                "initial",
            )

            # Этап 4: Интеллектуальная оценка и при необходимости расширенный поиск
            articles_combined_raw = await self._assess_and_expand_search(
                question,
                category,
                articles_initial_raw,
                law_files,
                chapters,
            )

            # ЭТАП 5: RAG2 - ТОЧНАЯ ВАЛИДАЦИЯ ВСЕХ НАЙДЕННЫХ СТАТЕЙ
            print("🎯 ЭТАП 5: RAG2 - Точная валидация всех найденных статей")
            rag2_results = self.rag2_validator.validate_articles(
                articles_combined_raw,
                question,
            )
            final_articles = rag2_results["final_articles"]

            # Сохраняем результаты RAG2 в лог
            self.search_log["rag2_validation"] = rag2_results
            print(f" 📊 RAG2 результаты:")
            print(f" ✅ Валидировано: {rag2_results['articles_validated']}")
            print(f" 🔧 Исправлено: {rag2_results['articles_corrected']}")
            print(f" ❌ Отклонено: {rag2_results['articles_rejected']}")

            # Сохраняем чистые, проверенные статьи в JSON перед генерацией
            self._save_final_articles_to_json(final_articles, question, category)

            # 🔥 НОВЫЙ ЭТАП 6: ГЕНЕРАЦИЯ ОТВЕТА НА ОСНОВЕ ВАЛИДИРОВАННЫХ СТАТЕЙ
            print("🧠 ЭТАП 6: Генерация ответа на основе валидированных статей")
            answer_generator = AnswerGenerator(self.project_root)

            generation_result = answer_generator.generate_answer_from_articles(
                question=question,
                articles_file_path=self.articles_file_path,
                category=category,
                search_log=self.search_log,
            )

            # Сохраняем результаты генерации в лог
            self.search_log["answer_generation"] = generation_result

            self.search_log["end_time"] = time.time()
            search_duration = self.search_log["end_time"] - self.search_log["start_time"]

            # 🔹 Подсчёт ресурсов за всё время поиска
            try:
                total_wall_time = time.time() - search_start_time

                if proc is not None and cpu_start is not None:
                    cpu_end = proc.cpu_times()
                    mem_end = proc.memory_info().rss

                    cpu_time_start = cpu_start.user + cpu_start.system
                    cpu_time_end = cpu_end.user + cpu_end.system
                    cpu_time_delta = max(0.0, cpu_time_end - cpu_time_start)

                    if total_wall_time > 0:
                        avg_cpu_cores = cpu_time_delta / total_wall_time
                        cpu_count = psutil.cpu_count(logical=True) or 1
                        avg_cpu_percent = min(
                            100.0,
                            max(0.0, avg_cpu_cores * 100.0 / cpu_count),
                        )
                    else:
                        avg_cpu_cores = 0.0
                        avg_cpu_percent = 0.0

                    max_rss_mb = max(mem_start or 0, mem_end) / (1024 * 1024)

                    resource_stats = {
                        "wall_time_sec": round(total_wall_time, 2),
                        "cpu_time_sec": round(cpu_time_delta, 2),
                        "avg_cpu_cores": round(avg_cpu_cores, 2),
                        "avg_cpu_percent": round(avg_cpu_percent, 1),
                        "memory_mb": round(max_rss_mb, 1),
                    }

                    print(
                        f"   🧮 РЕСУРСЫ ПОИСКА: время = {resource_stats['wall_time_sec']} c, "
                        f"CPU ≈ {resource_stats['avg_cpu_percent']:.1f}% "
                        f"({resource_stats['avg_cpu_cores']:.2f} ядер), "
                        f"RAM ≈ {resource_stats['memory_mb']:.1f} МБ"
                    )
                else:
                    print(
                        f"   🧮 РЕСУРСЫ ПОИСКА: замер CPU недоступен, "
                        f"время = {total_wall_time:.2f} c"
                    )
            except Exception as e:
                print(f"   ⚠️ Не удалось замерить ресурсы поиска: {e}")

            # Кладём статистику в search_log
            self.search_log["resource_stats"] = resource_stats

            # Сохраняем полный лог
            self._save_search_log(log_file_path)

            print(f" ✅ Поиск завершен: {len(final_articles)} статей, {search_duration:.2f}с")
            print(f" 📊 Лог поиска: {log_file_path}")

            return {
                "validation_decision": "sequential_search_complete",
                "reasoning": self.search_log["sufficiency_assessment"],
                "chapters_evaluation": self.search_log["chapters_processed"],
                "found_articles": final_articles,
                "search_log": self.search_log,
                "articles_file": str(self.articles_file_path),
                "log_file": str(log_file_path),
                "rag2_validation": rag2_results,      # результаты RAG2 в ответ
                "answer_generation": generation_result,  # результат генерации
            }

        except Exception as e:
            error_msg = f"Критическая ошибка поиска: {e}"
            print(f"   ❌ {error_msg}")
            import traceback
            traceback.print_exc()
            return self._create_error_result(error_msg, log_file_path)
    
    def _get_laws_for_category(self, category: str) -> List[str]:
        """Получает список законов для категории"""
        try:
            # Импортируем config из корня проекта
            import sys
            sys.path.insert(0, str(self.project_root))
            from config import config
            
            law_files = config.get_extended_laws_for_category(category)
            
            # Фильтруем только существующие файлы
            existing_law_files = []
            for law_file in law_files:
                law_path = self.laws_dir / law_file
                if law_path.exists():
                    existing_law_files.append(law_file)
                    print(f"   ✅ Файл найден: {law_file}")
                else:
                    print(f"   ⚠️ Файл закона не найден: {law_path}")
            
            self.search_log["laws_searched"] = existing_law_files
            
            if not existing_law_files:
                print(f"   🚨 Нет доступных файлов законов для категории '{category}'")
                return []
            
            return existing_law_files
            
        except Exception as e:
            print(f"   ⚠️ Ошибка получения законов: {e}")
            return []
    
    def _get_all_chapters_from_laws(self, law_files: List[str]) -> List[Dict]:
        """
        Получает все главы из законов.
        Для DOCX использует DocxIntelligentParser v3 с умной нарезкой.
        """
        all_chapters = []
        # Инициализируем наш новый парсер
        smart_parser = DocxIntelligentParser() 

        for law_file in law_files:
            try:
                law_path = self.laws_dir / law_file
                if not law_path.exists():
                    print(f" ⚠️ Файл не найден: {law_file}")
                    continue

                # === НОВАЯ ЛОГИКА ДЛЯ DOCX ===
                if law_file.lower().endswith('.docx'):
                    print(f" 📖 Иерархический парсинг DOCX: {law_file}")
                    try:
                        # 1. Строим дерево (Tree Building)
                        # Используем smart_parser, который умеет читать XML outlineLvl
                        tree_root = smart_parser.parse(str(law_path))
                        
                        # 2. Получаем умные чанки (Tree Chunking)
                        # Лимит 120k символов, сохраняем иерархию
                        chunks = smart_parser.chunk_tree(tree_root, max_chars=120000)
                        
                        # 3. Адаптируем под формат Searcher'а
                        # Searcher привык к ключам: 'law', 'chapter_title', 'content'
                        valid_chapters = []
                        for ch in chunks:
                            valid_chapters.append({
                                'law': law_file,
                                # Маппинг имен: chunk_title -> chapter_title
                                'chapter_title': ch['chunk_title'], # напр. "Глава IX > § 4"
                                'title': ch['chunk_title'],         # Дублируем для надежности
                                'content': ch['content'],
                                'type': ch['type'],
                                'level': ch['level']
                            })
                        
                        all_chapters.extend(valid_chapters)
                        print(f" -> Извлечено {len(valid_chapters)} структурных чанков (v3)")

                    except Exception as e:
                        print(f" ❌ Ошибка Smart Docx Parser для {law_file}: {e}")
                        # Fallback на старый метод чтения текста
                        self._read_law_fallback(law_path, law_file, all_chapters)
                
                # === ЛОГИКА ДЛЯ TXT ===
                else:
                    self._read_law_fallback(law_path, law_file, all_chapters)

            except Exception as e:
                print(f" ❌ Критическая ошибка обработки закона {law_file}: {e}")
                
        return all_chapters

    def _read_law_fallback(self, law_path: Path, law_file: str, all_chapters: List[Dict]):
        """
        Вспомогательный метод для чтения TXT или 'плохих' DOCX как плоского текста.
        Использует Regex-парсинг.
        """
        try:
            content = ""
            if law_file.lower().endswith('.docx'):
                # Читаем DOCX "тупо" параграф за параграфом
                try:
                    import docx
                    doc = docx.Document(law_path)
                    full_text = [p.text for p in doc.paragraphs]
                    content = "\n".join(full_text)
                    print(f" 📖 Прочитан DOCX файл: {law_file} ({len(content)} симв.)")
                except Exception as docx_error:
                    error_msg = str(docx_error)
                    if "word/word/" in error_msg or "settings.xml" in error_msg:
                        print(f" ⚠️ DOCX файл {law_file} имеет нестандартную структуру: {docx_error}")
                        print(f" ⏭️ Пропускаю файл, продолжаю обработку остальных законов")
                        return
                    else:
                        print(f" ⚠️ Ошибка чтения DOCX файла {law_file}: {docx_error}")
                        print(f" ⏭️ Пропускаю файл, продолжаю обработку остальных законов")
                        return
            else:
                # Читаем TXT с автоопределением кодировки
                try:
                    with open(law_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    with open(law_path, 'r', encoding='windows-1251') as f:
                        content = f.read()
                print(f" 📖 Прочитан TXT файл: {law_file} ({len(content)} симв.)")

            if not content:
                print(f" ⚠️ Файл пуст или не содержит текста: {law_file}")
                return

            # Используем старый Regex-экстрактор (он есть в классе)
            # Этот метод ищет "Глава X", "Статья Y" по тексту
            chapters_from_law = self._extract_all_chapters_from_law(content, law_file)
            all_chapters.extend(chapters_from_law)
            print(f" -> Извлечено глав (Regex): {len(chapters_from_law)}")

        except Exception as e:
            print(f" ❌ Ошибка в fallback-чтении {law_file}: {e}")

    def _extract_all_chapters_from_law(self, law_content: str, law_name: str) -> List[Dict]:
        """Извлекает все структурные единицы из закона с учетом типа документа"""
        doc_type = self._detect_document_type(law_name)
        chapters = []
        
        print(f"   🔍 Извлечение структур для {law_name} (тип: {doc_type})")
        
        if doc_type == "code":
            # 🔷 ДЛЯ КОДЕКСОВ - СТАРАЯ ЛОГИКА (оставляем как было)
            chapters = self._extract_chapters_from_code(law_content, law_name)
            
        elif doc_type == "plenum":
            # 📘 ДЛЯ ПЛЕНУМОВ - используем специальный парсер
            structures = self.simple_parser._parse_plenum_by_headings(law_content, law_name)
            chapters = self._convert_structures_to_chapters(structures, law_name)
            
        elif doc_type == "sublegal":
            # 📋 ДЛЯ ПОДЗАКОННЫХ АКТОВ - используем парсер статей
            structures = self.simple_parser._parse_sublegal_structure(law_content, law_name)
            chapters = self._convert_structures_to_chapters(structures, law_name)
            
        else:
            # 🔸 FALLBACK - используем общий парсер простых структур
            structures = self.simple_parser.parse_simple_structure(law_content, law_name)
            chapters = self._convert_structures_to_chapters(structures, law_name)
        
        print(f"   ✅ Извлечено структурных единиц: {len(chapters)}")
        return chapters

    # ДОБАВИТЬ НОВЫЕ ВСПОМОГАТЕЛЬНЫЕ МЕТОДЫ:
    def _extract_chapters_from_code(self, law_content: str, law_name: str) -> List[Dict]:
        """ИЗВЛЕЧЕНИЕ ГЛАВ ИЗ КОДЕКСОВ - СУЩЕСТВУЮЩАЯ ЛОГИКА"""
        chapters = []
        lines = law_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Ищем начало главы (можно добавить другие паттерны)
            if (line.startswith('ГЛАВА') or line.startswith('Глава') or 
                line.startswith('РАЗДЕЛ') or line.startswith('Раздел')):
                
                chapter_title = line
                chapter_start = i
                
                # Ищем конец главы
                chapter_end = len(lines)
                for j in range(i + 1, len(lines)):
                    next_line = lines[j].strip()
                    # Следующая глава или раздел
                    if ((next_line.startswith('ГЛАВА') or next_line.startswith('Глава') or
                         next_line.startswith('РАЗДЕЛ') or next_line.startswith('Раздел')) and
                        j > i + 5):  # Минимум 5 строк в главе
                        chapter_end = j
                        break
                
                # Извлекаем содержание главы
                chapter_content = '\n'.join(lines[chapter_start:chapter_end])
                
                chapters.append({
                    'law': law_name,
                    'title': chapter_title,
                    'content': chapter_content
                })
                
                i = chapter_end  # Переходим к следующей главе
            else:
                i += 1
        
        return chapters

    def _convert_structures_to_chapters(self, structures: List[Dict], law_name: str) -> List[Dict]:
        """Конвертирует структуры из SimpleLawParser в формат глав"""
        chapters = []

        for structure in structures:
            # Пропускаем преамбулы и слишком мелкие структуры
            if structure.get('type') == 'preamble':
                continue
                
            # Создаем главу в нужном формате
            chapter = {
                'law': law_name,
                'title': structure.get('title', 'Без названия'),
                'content': structure.get('content', ''),
                'structure_type': structure.get('type', 'unknown')  # сохраняем тип для диагностики
            }
            
            # Фильтруем слишком короткие содержания
            if len(chapter['content']) > 100:  # минимум 100 символов
                chapters.append(chapter)
            else:
                print(f"   ⚠️  Пропущена короткая структура: {chapter['title'][:50]}...")
        
        return chapters    
    
    def _find_relevant_chapters(self, question: str, law_files: List[str], category: str) -> List[Dict]:
        """Находит релевантные главы через навигатор"""
        if self.navigator:
            try:
                print(f"   🔍 Использую навигатор для поиска глав...")
                chapters = self.navigator.find_relevant_chapters(question, law_files, max_chapters=5)
            
                # 🔥 ДИАГНОСТИКА: какие типы документов найдены
                doc_types = {}
                for chapter in chapters:
                    law_name = chapter.get('law', 'Unknown')
                    doc_type = self._detect_document_type(law_name)
                    doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
                    print(f"      📖 {law_name} -> {doc_type}: {chapter.get('title', 'No title')[:50]}...")
                    
                print(f"   📊 Найдено структур по типам: {doc_types}")
                
                # Обогащаем главы реальным содержанием из файлов
                enriched_chapters = []
                for chapter in chapters:
                    enriched_chapter = self._enrich_chapter_content(chapter)
                    enriched_chapters.append(enriched_chapter)
                
                return enriched_chapters
                
            except Exception as e:
                print(f"   ⚠️ Ошибка навигатора: {e}")
                import traceback
                traceback.print_exc()
        
        # Если навигатора нет - возвращаем пустой список
        print(f"   ❌ Навигатор недоступен")
        return []
    
    async def _analyze_chapters_phase(
        self,
        question: str,
        chapters: List[Dict],
        category: str,
        phase: str,
    ) -> List[Dict]:
        """Асинхронно анализирует главы в указанной фазе поиска"""
        print(f" 🔍 Фаза {phase}: анализ {len(chapters)} глав")
        phase_articles = []

        # Фильтруем главы для обработка
        chapters_to_process = []
        for chapter in chapters:
            title = chapter.get('title') or chapter.get('chapter_title') or 'Unknown'
            chapter_id = f"{chapter.get('law', 'Unknown')}_{title}" 

            if phase == "initial" and chapter_id in self.processed_chapters:
                print(f" ⏭️ Глава уже обработана (initial): {chapter.get('title', 'No title')[:50]}...")
                continue

            # Помечаем главу как обработанную хотя бы в одной фазе
            self.processed_chapters.add(chapter_id)
            chapters_to_process.append(chapter)

        if not chapters_to_process:
            return phase_articles

        # Создаем асинхронные задачи для анализа глав
        tasks = []
        for i, chapter in enumerate(chapters_to_process, 1):
            print(f" 📖 {phase} - Глава {i}/{len(chapters_to_process)}: {chapter.get('title', 'No title')[:50]}...")
            
            task = self._analyze_single_chapter_async(
                question,
                chapter,
                category,
                phase,
                i,
            )
            tasks.append(task)

        # Запускаем задачи параллельно с ограничением одновременных запросов (50 одновременно)
        semaphore = asyncio.Semaphore(50)
        
        async def run_with_semaphore(task):
            async with semaphore:
                return await task
        
        results = await asyncio.gather(*[run_with_semaphore(task) for task in tasks], return_exceptions=True)

        # Обрабатываем результаты
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                print(f"   ❌ Ошибка при анализе главы {i+1}: {result}")
                chapter_result = {
                    "chapter_title": chapters_to_process[i].get('title', 'No title'),
                    "law": chapters_to_process[i].get('law', 'Unknown'),
                    "phase": phase,
                    "articles_found": [],
                    "analysis_quality": "error",
                    "error": str(result)
                }
            else:
                chapter_result = result

            self.search_log["chapters_processed"].append(chapter_result)
            self.search_log["total_chapters_analyzed"] += 1

            # Сохраняем найденные статьи
            if chapter_result.get("articles_found"):
                new_articles = chapter_result["articles_found"]
                phase_articles.extend(new_articles)
                await self._save_articles_to_file_async(new_articles, chapters_to_process[i], phase)
                print(f" ✅ Найдено статей: {len(new_articles)}")
            else:
                print(f" ℹ️ Статей не найдено")

        return phase_articles            
    
    def _enrich_chapter_content(self, chapter: Dict) -> Dict:
        """Обогащает главу ПОЛНЫМ содержанием включая все дочерние структуры"""
        try:
            law_file = self.laws_dir / chapter['law']
        
            if law_file.exists():
                # Используем навигатор для получения полного содержания главы
                if self.navigator:
                    full_content = self.navigator.get_full_chapter_content(chapter)
                    if full_content and len(full_content) > len(chapter.get('content', '')):
                        chapter['content'] = full_content
                        print(f"   📄 Загружено ПОЛНОЕ содержание главы: {len(full_content)} символов")
                    else:
                        # Fallback: используем старую логику
                        with open(law_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                    
                        chapter_content = self._extract_chapter_from_law(content, chapter['title'])
                        if chapter_content:
                            chapter['content'] = chapter_content
                            print(f"   📄 Загружено обычное содержание: {len(chapter_content)} символов")
                else:
                    print(f"   ⚠️  Навигатор недоступен, использую обычное содержание")
            else:
                print(f"   ⚠️  Файл закона не найден: {law_file}")
            
        except Exception as e:
            print(f"   ⚠️  Ошибка загрузки содержания: {e}")
    
        return chapter
    
    def _extract_chapter_from_law(self, law_content: str, chapter_title: str) -> str:
        """Извлекает содержание главы из текста закона"""
        # Упрощенный поиск главы по заголовку
        lines = law_content.split('\n')
        chapter_start = -1
        
        # Ищем начало главы
        for i, line in enumerate(lines):
            if chapter_title.lower() in line.lower():
                chapter_start = i
                break
        
        if chapter_start == -1:
            return None
        
        # Ищем конец главы (следующую главу или раздел)
        next_chapter_patterns = [
            'ГЛАВА', 'РАЗДЕЛ', 'СТАТЬЯ', 'Глава', 'Раздел', 'Статья'
        ]
        
        chapter_end = len(lines)
        for i in range(chapter_start + 1, len(lines)):
            line = lines[i].strip()
            if any(pattern in line for pattern in next_chapter_patterns) and len(line) < 200:
                # Проверяем, что это действительно начало новой структурной единицы
                if i > chapter_start + 5:  # Минимум 5 строк в главе
                    chapter_end = i
                    break
        
        return '\n'.join(lines[chapter_start:chapter_end])
    
    
    def _build_chunk_aware_prompt(self, question: str, chapter: Dict, category: str, phase: str) -> str:
        """
        Строит промпт. Если это подчанк (Параграф/Раздел), добавляет контекстную справку.
        """
        base_prompt = self._build_real_search_prompt(question, chapter, category, phase)

        # Если это чанк (структурный блок), добавляем навигационную помощь
        if '[Чанк' in chapter.get('title', ''):
            # Извлекаем "чистое" название главы и номер чанка для ясности
            # Пример title: "Глава X. Банкротство [Чанк 2/5]"
            clean_title = chapter['title'].split('[')[0].strip()
            
            chunk_info = f"""
            ℹ️ КОНТЕКСТНАЯ СПРАВКА:
            Вы анализируете фрагмент из документа: "{chapter.get('law', 'Закон')}"
            Раздел/Глава: "{clean_title}"
            
            ЭТОТ ФРАГМЕНТ ЯВЛЯЕТСЯ ЛОГИЧЕСКИ ЦЕЛОСТНЫМ БЛОКОМ (например, отдельным Параграфом или группой статей).
            Анализируйте его как самостоятельный текст. Ищите статьи ВНУТРИ этого блока.
            """
            return chunk_info + "\n\n" + base_prompt
            
        return base_prompt
    
    def _analyze_single_chapter(self, question: str, chapter: Dict, category: str, phase: str, chapter_index: int) -> Dict:
        """Анализ одной главы с изолированным контекстом"""
        
        law_name = chapter.get('law', 'Unknown')
        chapter_title = chapter.get('title', 'No title')
        chapter_content = chapter.get('content', '')
        
        if not chapter_content:
            print(f"   ⚠️  Пустое содержание главы, пропускаем")
            return {
                "chapter_title": chapter_title,
                "law": law_name,
                "phase": phase,
                "articles_found": [],
                "analysis_quality": "empty_content",
                "error": "Пустое содержание главы"
            }
        
        # Промпт для буквального поиска в реальном законе
        prompt = self._build_chunk_aware_prompt(question, chapter, category, phase)
        
        try:
            # Изолированный запрос для этой главы
            messages = [{"role": "user", "content": prompt}]
            response = api_manager.deepseek_completion(
                messages=messages,
                temperature=0.4,
                timeout=90,
                response_format={"type": "json_object"}
            )
            
            # Парсим результат
            chapter_analysis = self._parse_chapter_response(response, chapter)
            chapter_analysis["phase"] = phase
            chapter_analysis["chapter_index"] = chapter_index
            chapter_analysis["law"] = law_name
            chapter_analysis["chapter_title"] = chapter_title
            
            # Логируем процесс поиска
            search_step = {
                "phase": phase,
                "chapter_index": chapter_index,
                "chapter_title": chapter_title,
                "law": law_name,
                "content_length": len(chapter_content),
                "response_length": len(response),
                "articles_found_count": len(chapter_analysis.get("articles_found", [])),
                "analysis_quality": chapter_analysis.get("analysis_quality", "unknown"),
                "timestamp": time.time()
            }
            self.search_log["search_sequence"].append(search_step)
            
            return chapter_analysis
            
        except Exception as e:
            print(f"   ❌ Ошибка анализа главы {chapter_index}: {e}")
            return {
                "chapter_title": chapter_title,
                "law": law_name,
                "phase": phase,
                "articles_found": [],
                "analysis_quality": "error",
                "error": str(e)
            }
    
    async def _analyze_single_chapter_async(
        self, 
        question: str, 
        chapter: Dict, 
        category: str, 
        phase: str, 
        chapter_index: int
    ) -> Dict:
        """Асинхронная версия анализа одной главы с поддержкой чанков"""
        
        # 🔥 ДОБАВЛЯЕМ: Проверка наличия поля law в главе
        if not chapter.get('law'):
            print(f"   🚨 КРИТИЧЕСКАЯ ОШИБКА: Глава '{chapter.get('title', 'Unknown')}' не имеет поля 'law'!")
            print(f"   Полная глава для отладки: {chapter}")
            return {
                "chapter_title": chapter.get('title', 'Unknown'),
                "law": "ERROR_NO_LAW_FIELD",
                "phase": phase,
                "articles_found": [],
                "analysis_quality": "error_no_law",
                "error": "Глава не содержит поля 'law'"
            }        
        
        # --- НАЧАЛО ИЗМЕНЕНИЙ (Защита: если нет текста, читаем с диска) ---
        if not chapter.get('content'):
            try:
                # Если RAG1 или Расширенный поиск не передали текст, читаем его с диска
                print(f"   📖 Читаю текст главы с диска: {chapter.get('title')}")
                # ВНИМАНИЕ: Тут может потребоваться адаптация под новый парсер, 
                # если self.parser - это старый SimpleLawParser.
                # Для совместимости пока пробуем читать старым методом или просто пропускаем,
                # так как новый HierarchicalDocxParser должен сразу отдавать контент.
                
                # content = self.parser.get_chapter_content(chapter.get('law'), chapter.get('title'))
                # chapter['content'] = content
                
                # Если контента нет - это критично, но для безопасности просто вернем ошибку
                raise ValueError("Контент главы пуст и не может быть загружен")

            except Exception as e:
                print(f"   ❌ Ошибка чтения главы с диска: {e}")
                return {
                    "chapter_title": chapter.get('title', 'Unknown'),
                    "law": chapter.get('law', 'Unknown'),
                    "phase": phase,
                    "articles_found": [],
                    "analysis_quality": "error_no_content",
                    "error": str(e)
                }
        # --- КОНЕЦ ИЗМЕНЕНИЙ ---

        # ТЕПЕРЬ ВМЕСТО self.chunk_manager.prepare_chapter_analysis(chapter)
        # Мы доверяем новому парсеру, который уже отдал нам чанк нужного размера.
        # Просто оборачиваем его в список для совместимости с циклом ниже.
        
        analysis_units = [{
            'type': chapter.get('type', 'chapter_chunk'),
            'title': chapter.get('title', 'No title'),
            'content': chapter.get('content', ''),
            'chunk_index': 0,          # Это всегда 0, так как чанк уже готов
            'total_chunks': 1,         # Всего 1
            'law': chapter.get('law', ''),
            'original_chapter': chapter.get('title', 'No title')
        }]
        
        # Старый вызов удаляем или комментируем:
        # analysis_units = self.chunk_manager.prepare_chapter_analysis(chapter)
        
        all_articles = []
        chunk_results = []
        
        # ... (далее идет цикл for unit_index, unit in enumerate(analysis_units): ... )
        
        # Анализируем каждый чанк отдельно
        for unit_index, unit in enumerate(analysis_units):
            print(f"   📖 Анализ {unit['type']}: {unit['title']}")
            
            try:
                # Создаем виртуальную "главу" для этого чанкf
                virtual_chapter = {
                    'law': unit.get('law', chapter.get('law', '')),
                    'title': unit['title'],
                    'content': unit['content']
                }
                
                # Анализируем чанк как отдельную главу
                result = await asyncio.get_event_loop().run_in_executor(
                    None, 
                    self._analyze_single_chapter_sync,
                    question, virtual_chapter, category, f"{phase}_chunk_{unit_index}", chapter_index
                )
                
                chunk_results.append({
                    'chunk_title': unit['title'],
                    'articles_found': result.get('articles_found', []),
                    'analysis_quality': result.get('analysis_quality', 'unknown'),
                    'chunk_index': unit['chunk_index'],
                    'total_chunks': unit['total_chunks']
                })
                
                # Добавляем найденные статьи с информацией о чанке
                for article in result.get('articles_found', []):
                    article['source_chunk'] = f"{unit['chunk_index'] + 1}/{unit['total_chunks']}"
                    all_articles.append(article)
                
                print(f"   ✅ Чанк {unit_index + 1}/{len(analysis_units)} обработан: {len(result.get('articles_found', []))} статей")
                
            except Exception as e:
                print(f"   ❌ Ошибка анализа чанка {unit_index + 1}: {e}")
                chunk_results.append({
                    'chunk_title': unit['title'],
                    'articles_found': [],
                    'analysis_quality': 'error',
                    'error': str(e)
                })
        
        # Объединяем результаты
        final_result = {
            "chapter_title": chapter.get('title', 'No title'),
            "law": chapter.get('law', ''),
            "phase": phase,
            "articles_found": all_articles,
            "analysis_quality": "multi_chunk" if len(analysis_units) > 1 else "single_chunk",
            "chunk_analysis": chunk_results,
            "total_chunks": len(analysis_units),
            "chunks_processed": len([r for r in chunk_results if r.get('articles_found')])
        }
        
        print(f"   📊 Итоги по главе: {len(all_articles)} статей из {len(analysis_units)} чанков")
        return final_result
    
    def _analyze_single_chapter_sync(
        self,
        question: str, 
        chapter: Dict, 
        category: str, 
        phase: str, 
        chapter_index: int
    ) -> Dict:
        """Синхронная обертка для анализа одной главы"""
        return self._analyze_single_chapter(question, chapter, category, phase, chapter_index)

    async def _save_articles_to_file_async(self, articles: List[Dict], chapter: Dict, phase: str):
        """Асинхронная версия сохранения статей в файл"""
        await asyncio.get_event_loop().run_in_executor(
            None, self._save_articles_to_file, articles, chapter, phase
        )    
    
    def _detect_document_type(self, law_name: str) -> str:
        """
        Грубое определение типа документа по имени файла:
        - code     — кодексы, КоАП, ПДД и т.п.
        - plenum   — Пленумы
        - sublegal — подзаконные акты (постановления и т.п.)
        """
        if not law_name:
            return "unknown"

        name = law_name.lower()

        # Пленумы Верховного Суда и т.п.
        if "plenum" in name:
            return "plenum"

        # Подзаконные акты (постановления правительства и т.п.)
        if "postanov" in name or "postanovlen" in name:
            return "sublegal"

        # Всё остальное по умолчанию считаем кодексами/законами
        return "code"
    
    def _get_search_hint(self, question: str, category: str) -> str:
        """Получает умный совет для текущей категории и вопроса"""
        if not hasattr(self, 'search_hints') or not self.search_hints:
            return ""
            
        # Получаем конфиг для категории
        category_hints = self.search_hints.get(category)
        if not category_hints:
            return ""
            
        # Если в конфиге есть общий текст подсказки (старый формат)
        if isinstance(category_hints, dict) and "hint_text" in category_hints:
             return category_hints["hint_text"]
             
        # Если в конфиге есть список триггеров (новый формат)
        # (Предполагаем структуру: {"triggers": [...], "hint": "..."})
        if isinstance(category_hints, dict) and "triggers" in category_hints:
            question_lower = question.lower()
            for trigger in category_hints["triggers"]:
                if trigger.lower() in question_lower:
                    return category_hints.get("hint", "")
                    
        return ""
    
    
    
    def _build_real_search_prompt(self, question: str, chapter: Dict, category: str, phase: str) -> str:
        """Строит промпт для поиска: ЖЕСТКИЙ, ТОЧНЫЙ, С ПОДДЕРЖКОЙ СОВЕТОВ"""
        
        law_name = chapter.get('law', 'Unknown')
        chapter_title = chapter.get('title', 'No title')
        chapter_content = chapter.get('content', '')
        doc_type = self._detect_document_type(law_name) # Убедитесь, что этот метод есть

        current_law_name = law_name  # Будем использовать в нескольких местах       
        
        # 1. Получаем "Умный совет" для ИИ (если есть)
        search_hint = self._get_search_hint(question, category)

        prompt_parts = []
        
        # --- СИСТЕМНАЯ РОЛЬ ---
        prompt_parts.append("РОЛЬ: Ты — робот-юрист. Твоя задача — НЕ думать, а ИСКАТЬ и ЦИТИРОВАТЬ.")
        prompt_parts.append("Твоя цель: Извлечь из текста ниже цитаты, отвечающие на вопрос пользователя.")
        prompt_parts.append("")
        
        # --- КОНТЕКСТ ---
        prompt_parts.append(f"ВОПРОС: {question}")
        prompt_parts.append(f"ДОКУМЕНТ: {law_name} ({chapter_title})")
        
        # --- ВНЕДРЕНИЕ УМНОГО СОВЕТА (В самый видный блок) ---
        if search_hint:
            prompt_parts.append("")
            prompt_parts.append("🔴 ОСОБАЯ ИНСТРУКЦИЯ (СОВЕТ):")
            prompt_parts.append(search_hint)
            prompt_parts.append("Если совет применим к тексту главы — СЛЕДУЙ ЕМУ ПРИОРИТЕТНО.")
            prompt_parts.append("Если в тексте нет того, о чем говорится в совете — ИГНОРИРУЙ СОВЕТ.")
        
        # --- ТЕКСТ ЗАКОНА ---
        prompt_parts.append("")
        prompt_parts.append("--- ТЕКСТ ДЛЯ АНАЛИЗА (НАЧАЛО) ---")
        prompt_parts.append(chapter_content)
        prompt_parts.append("--- ТЕКСТ ДЛЯ АНАЛИЗА (КОНЕЦ) ---")
        prompt_parts.append("")

        # --- ИНСТРУКЦИИ ПО ИЗВЛЕЧЕНИЮ (Защита от галлюцинаций) ---
        prompt_parts.append("ПРАВИЛА ИЗВЛЕЧЕНИЯ:")
        
        if doc_type == "code":
            # Для кодексов требуем жесткий формат номера
            prompt_parts.append("1. ПОЛЕ 'article': Строгий формат 'Статья <НОМЕР>'.")
            prompt_parts.append("   - ✅ Правильно: 'Статья 18', 'Статья 213.6', 'Статья 15.1'.")
            prompt_parts.append("   - ❌ Неправильно: 'Статья 18 (О правах)', 'ст. 18', '18'.")        
        
        elif doc_type in ["plenum", "sublegal", "decree"]: 
            # Для Пленумов, Постановлений, Указов - ищем Пункты
            prompt_parts.append("1. В поле 'article' пиши номер ПУНКТА или АБЗАЦА (пример: 'Пункт 1', 'п. 2.1', 'Абзац 3').")
            prompt_parts.append("2. Внимание! В этом документе может не быть 'Статей'. Считай каждый нумерованный пункт (1., 2.) отдельной нормой.")

        else:
            # Fallback для неизвестных типов (договоры, обзоры и т.д.)
            prompt_parts.append("1. В поле 'article' пиши наиболее точный идентификатор: 'Статья X', 'Пункт Y', или 'Раздел Z'.")
            prompt_parts.append("2. Если явной нумерации нет, используй название раздела или первые 3-4 слова абзаца как идентификатор.")            
        
        prompt_parts.append("3. ТОЧНОСТЬ ЦИТИРОВАНИЯ: В поле 'content' копируй текст БУКВАЛЬНО, символ в символ. Не перефразируй, не сокращай. Если сомневаешься — лучше скопировать больше текста, чем исказить смысл.")
        prompt_parts.append("4. Если статья отвечает на вопрос лишь частично — копируй её всё равно.")
        prompt_parts.append("5. Если в тексте НЕТ подходящих статей — верни пустой список found_articles: []. НЕ ВЫДУМЫВАЙ.")

    # 🔥 ДОБАВЛЯЕМ: Правило для поля 'law'
        prompt_parts.append(f"6. В поле 'law' ВСЕГДА указывай точное название файла закона: '{current_law_name}'. Не придумывай свои варианты!")        
        
        
        # --- ФОРМАТ ОТВЕТА ---
        prompt_parts.append("")
        prompt_parts.append("🔥 ТВОЙ ОТВЕТ: ВАЛИДНЫЙ JSON БЕЗ ЛЮБЫХ ДРУГИХ ТЕКСТОВ.")
        prompt_parts.append("   • Начинается с '{', заканчивается '}'.")
        prompt_parts.append("   • Никаких ```json, пояснений, 'Вот результат:' или markdown.")
        prompt_parts.append("   • Если статей нет: {\"articles_found\": [], \"analysis_quality\": \"none\"}")
        prompt_parts.append("")
        prompt_parts.append("ФОРМАТ ОТВЕТА (JSON):")
        prompt_parts.append("{")
        prompt_parts.append('  "analysis_quality": "high",')
        prompt_parts.append('  "articles_found": [')
        prompt_parts.append('    {')        
        
        prompt_parts.append(f'      "law": "{current_law_name}",  # ИСПОЛЬЗУЙ ЭТО ТОЧНОЕ НАЗВАНИЕ')    
        
        # Пример зависит от типа документа, чтобы ИИ понял паттерн
        if doc_type == "code":
            prompt_parts.append('      "article": "Статья 123 (ТОЛЬКО НОМЕР!)",')
        else:
            prompt_parts.append('      "article": "Пункт 10",')
        prompt_parts.append('      "content": "Точная цитата из текста...",')
        prompt_parts.append('      "relevance_reasoning": "Почему это подходит..."')
        prompt_parts.append('    }')
        prompt_parts.append('  ]')
        prompt_parts.append("}")
        
        return "\n".join(prompt_parts)
    
    async def _assess_and_expand_search(self, question: str, category: str, current_articles: List[Dict], 
                                law_files: List[str], original_chapters: List[Dict]) -> List[Dict]:
        """Оценивает достаточность статей с помощью ИИ и запускает расширенный поиск при необходимости"""
        
        print(f"   📊 Интеллектуальная оценка результатов: {len(current_articles)} статей")
        
        # Если статей вообще нет - сразу расширенный поиск
        if len(current_articles) == 0:
            assessment = "Статьи не найдены. Запускаем расширенный поиск."
            self.search_log["sufficiency_assessment"] = assessment
            self.search_log["expansion_triggered"] = True
            print(f"   🔍 {assessment}")
            expanded_articles = await self._perform_expanded_search(question, category, law_files, current_articles)
            return current_articles + expanded_articles
        
        # Интеллектуальная оценка достаточности через ИИ
        assessment_result = self._assess_articles_sufficiency(question, current_articles)
        
        if assessment_result["sufficient"]:
            assessment = f"Статьи достаточны для ответа: {assessment_result['reasoning']}"
            self.search_log["sufficiency_assessment"] = assessment
            print(f"   ✅ {assessment}")
            return current_articles
        else:
            assessment = f"Статьи недостаточны: {assessment_result['reasoning']}. Запускаем расширенный поиск."
            self.search_log["sufficiency_assessment"] = assessment
            self.search_log["expansion_triggered"] = True
            print(f"   🔍 {assessment}")
            expanded_articles = await self._perform_expanded_search(question, category, law_files, current_articles)
            return current_articles + expanded_articles

    def _assess_articles_sufficiency(self, question: str, articles: List[Dict]) -> Dict:
        """Оценивает, достаточно ли статей для полного ответа на вопрос"""
        
        # Подготавливаем информацию о статьях для промпта
        articles_info = []
        for i, article in enumerate(articles, 1):
            articles_info.append(f"{i}. {article.get('article', 'Без названия')}: {article.get('relevance_reasoning', '')}")
        
        articles_text = "\n".join(articles_info)
        
        prompt = f"""
🔍 ЗАДАЧА: ОЦЕНКА ДОСТАТОЧНОСТИ НАЙДЕННЫХ СТАТЕЙ

ВОПРОС ПОЛЬЗОВАТЕЛЯ: {question}

НАЙДЕННЫЕ СТАТЬИ:
{articles_text}

ПРОАНИЛИЗИРУЙ:
1. Полнота охвата - покрывают ли статьи все аспекты вопроса?
2. Конкретика - содержат ли статьи конкретные правовые нормы, а не общие положения?
3. Применимость - достаточно ли статей для практического ответа на вопрос?

КРИТЕРИИ НЕДОСТАТОЧНОСТИ:
- ❌ Есть пробелы в регулировании (не хватает статей по ключевым аспектам)
- ❌ Статьи слишком общие, без конкретных норм
- ❌ Не хватает процедурных моментов (сроки, порядок, документы)
- ❌ Вопрос требует ссылок на несколько взаимодополняющих статей, а найдена только одна

КРИТЕРИИ ДОСТАТОЧНОСТИ:
- ✅ Статьи покрывают все основные аспекты вопроса
- ✅ Содержат конкретные правовые нормы и процедуры
- ✅ Позволяют дать практический совет

ВЕРНИ ОТВЕТ В ФОРМАТЕ JSON:
{{
    "sufficient": true/false,
    "reasoning": "Развернутое обоснование на русском языке",
    "missing_aspects": ["Какие аспекты вопроса не покрыты", "Что еще нужно найти"]
}}
"""
        
        try:
            messages = [{"role": "user", "content": prompt}]
            response = api_manager.deepseek_completion(
                messages=messages,
                temperature=0.1,
                timeout=60,
                response_format={"type": "json_object"}
            )
            
            # Парсим ответ
            result = self._parse_sufficiency_response(response)
            return result
            
        except Exception as e:
            print(f"   ❌ Ошибка оценки достаточности: {e}")
            # В случае ошибки считаем, что недостаточно (перестраховываемся)
            return {
                "sufficient": False,
                "reasoning": f"Ошибка оценки: {str(e)}",
                "missing_aspects": ["Не удалось оценить полноту покрытия"]
            }

    def _parse_sufficiency_response(self, response: str) -> Dict:
        """Парсит ответ ИИ на оценку достаточности"""
        
        try:
            # Ищем JSON в ответе
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0).strip()
                result = json.loads(json_str)
                
                # Валидация структуры
                if "sufficient" not in result:
                    result["sufficient"] = False
                if "reasoning" not in result:
                    result["reasoning"] = "Не указано обоснование"
                if "missing_aspects" not in result:
                    result["missing_aspects"] = []
                    
                return result
            else:
                return {
                    "sufficient": False,
                    "reasoning": "Не удалось распарсить ответ ИИ",
                    "missing_aspects": ["Ошибка парсинга оценки"]
                }
                
        except Exception as e:
            return {
                "sufficient": False,
                "reasoning": f"Ошибка парсинга: {str(e)}",
                "missing_aspects": ["Ошибка обработки оценки"]
            }

    async def _perform_expanded_search(
        self,
        question: str,
        category: str,
        law_files: List[str],
        existing_articles: List[Dict],
    ) -> List[Dict]:
        """Выполняет расширенный поиск по всем законам категории"""
        print(f" 🚀 ЗАПУСК РАСШИРЕННОГО ПОИСКА")
        print(f" 📚 Анализируем все законы категории: {len(law_files)} файлов")

        # Получаем ВСЕ главы из ВСЕХ законов категории
        all_chapters = self._get_all_chapters_from_laws(law_files)
        print(f" 📖 Всего глав в категории (для expanded): {len(all_chapters)}")

        if not all_chapters:
            print(f" ℹ️ В законах категории не найдено глав для расширенного поиска")
            return []

        # Пытаемся отфильтровать уже обработанные главы (initial)
        new_chapters = []
        for chapter in all_chapters:
            title = chapter.get('title') or chapter.get('chapter_title') or 'Unknown'
            chapter_id = f"{chapter.get('law', 'Unknown')}_{title}"
            
            if chapter_id not in self.processed_chapters:
                new_chapters.append(chapter)

        print(f" 🔍 Новых глав для анализа (после фильтра): {len(new_chapters)}")

        # КРИТИЧЕСКОЕ ИЗМЕНЕНИЕ:
        # если после фильтрации не осталось новых глав, но all_chapters не пуст,
        # принудительно анализируем ВСЕ главы ещё раз в фазе expanded
        if not new_chapters and all_chapters:
            print(
                " ⚠️ Новых глав нет, но принудительно запускаем анализ всех глав "
                "в расширенном поиске (expanded)"
            )
            new_chapters = all_chapters

        if not new_chapters:
            # Сюда попадём только если all_chapters был пуст
            print(f" ℹ️ Нет глав для расширенного поиска")
            return []

        # Анализируем главы в фазе expanded
        expanded_articles = await self._analyze_chapters_phase(
            question,
            new_chapters,
            category,
            "expanded",
        )
        print(f" ✅ Расширенный поиск завершен: найдено {len(expanded_articles)} новых статей")
        return expanded_articles

    def _save_articles_to_file(self, articles: List[Dict], chapter: Dict, phase: str):
        """Сохраняет найденные статьи во внешний файл"""
        try:
            with open(self.articles_file_path, "a", encoding="utf-8") as f:
                f.write(f"\n{'='*80}\n")
                f.write(f"# ФАЗА: {phase}\n")
                f.write(f"# ЗАКОН: {chapter.get('law', 'Unknown')}\n")
                f.write(f"# ГЛАВА: {chapter.get('title', 'Unknown')}\n")
                f.write(f"# ВРЕМЯ: {time.ctime()}\n")
                f.write(f"{'='*80}\n\n")
                
                for i, article in enumerate(articles, 1):
                    f.write(f"СТАТЬЯ {i}:\n")
                    f.write(f"Название: {article.get('article', 'Unknown')}\n")
                    f.write(f"Релевантность: {article.get('relevance_reasoning', 'Не указана')}\n")
                    f.write(f"Точное совпадение: {article.get('exact_match', False)}\n")
                    f.write("-" * 60 + "\n")
                    f.write(article.get('content', '') + "\n")
                    f.write("\n" + "=" * 80 + "\n\n")
                    
        except Exception as e:
            print(f"   ❌ Ошибка сохранения статей: {e}")
    
    def _save_search_log(self, log_file_path: Path):
        """Сохраняет полный лог поиска в JSON файл"""
        try:
            with open(log_file_path, "w", encoding="utf-8") as f:
                json.dump(self.search_log, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"   ❌ Ошибка сохранения лога: {e}")
    
    def _create_error_result(self, error_msg: str, log_file_path: Path) -> Dict:
        """Создает результат с ошибкой"""
        self.search_log["end_time"] = time.time()
        self.search_log["sufficiency_assessment"] = f"Ошибка: {error_msg}"
        
        try:
            self._save_search_log(log_file_path)
        except:
            pass
        
        return {
            "validation_decision": "search_failed",
            "reasoning": error_msg,
            "chapters_evaluation": [],
            "found_articles": [],
            "search_log": self.search_log,
            "articles_file": str(self.articles_file_path) if self.articles_file_path else "N/A",
            "log_file": str(log_file_path)
        }

    def _parse_chapter_response(self, response: str, chapter: Dict) -> Dict:
        """Парсит ответ ИИ и извлекает найденные статьи с улучшенной обработкой ошибок"""

        cleaned_response = response.strip()
        print(f"        📨 Ответ ИИ ({len(cleaned_response)} символов): {cleaned_response[:200]}...")

        # 🔥 ДОБАВИТЬ ЭТО ПЕРЕД СУЩЕСТВУЮЩИМ КОДОМ:
        # Быстрая попытка парсинга всего ответа как JSON
        try:
            # Пробуем распарсить весь ответ как JSON
            result = json.loads(cleaned_response)
            print(f"        ✅ Успешный прямой парсинг всего ответа")
            # Приводим к нужной структуре
            if "articles_found" not in result:
                result["articles_found"] = []
            if "analysis_quality" not in result:
                result["analysis_quality"] = "unknown"
            return result

        except json.JSONDecodeError:
            pass  # Продолжаем обычную обработку

        # ДОБАВЛЕНО: Сохраняем оригинальный ответ для отладки
        original_response = cleaned_response

        try:
            # Пытаемся найти JSON в ответе
            json_str = None

            # Ищем в блоках ```json
            json_match = re.search(
                r'```json\s*(.*?)\s*```',
                cleaned_response,
                re.DOTALL)
            if json_match:
                json_str = json_match.group(1).strip()

            # Ищем любой JSON объект
            if not json_str:
                json_match = re.search(r'\{.*\}', cleaned_response, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0).strip()

            # Берем весь ответ как JSON
            if not json_str and cleaned_response.startswith(
                    '{') and cleaned_response.endswith('}'):
                json_str = cleaned_response

            if json_str:
                print(f"        🔍 Найден JSON для парсинга")

                # Очистка JSON от возможных проблемных символов
                json_str = self._clean_json_string(json_str)

                try:
                    result = json.loads(json_str)
                except json.JSONDecodeError as e:
                    print(f"        ⚠️  Первая попытка парсинга не удалась: {e}")
                    
                    # 🔥 ДОБАВЛЕНО: Сохраняем ответ для отладки
                    debug_dir = self.project_root / "debug" / "parse_errors"
                    debug_dir.mkdir(parents=True, exist_ok=True)
                    timestamp = int(time.time())
                    law_name = chapter.get('law', 'unknown').replace('.', '_')
                    chapter_title = chapter.get('title', 'unknown').replace('/', '_').replace('\\', '_')
                    safe_chapter_title = re.sub(r'[^\w\s-]', '', chapter_title)[:50]
                    filename = f"{timestamp}_{law_name}_{safe_chapter_title}.txt"
                    filepath = debug_dir / filename
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(f"Глава: {chapter.get('title', 'unknown')}\n")
                        f.write(f"Закон: {chapter.get('law', 'unknown')}\n")
                        f.write(f"Ответ ИИ (первые 10000 символов):\n{response[:10000]}\n")
                        f.write(f"Извлеченная JSON строка:\n{json_str[:2000]}\n")
                        f.write(f"Ошибка: {e}\n")
                    print(f"        🔍 Ответ сохранен в {filepath}")
                    
                    # Пытаемся восстановить JSON
                    repaired_json = self._repair_json_string(json_str)
                    if repaired_json:
                        print(f"        🔧 Пытаюсь восстановить JSON...")
                        result = json.loads(repaired_json)
                    else:
                        raise e

                # Валидируем структуру
                if "articles_found" not in result:
                    result["articles_found"] = []
                if "analysis_quality" not in result:
                    result["analysis_quality"] = "unknown"

                # ИЗМЕНЕНИЕ: Убираем агрессивную фильтрацию по качеству анализа
                if result.get("analysis_quality") in [
                        "none", "low", "no_articles"]:
                    print(
                        f"        ⚠️  ИИ сообщил о низком качестве анализа, но передаем статьи в RAG2 для проверки")

                # ИЗМЕНЕНИЕ: Убираем предварительную валидацию статей - пусть
                # RAG2 сам решает
                valid_articles = []
                for article in result["articles_found"]:
                    if isinstance(article, dict) and article.get('content'):
                        # Гарантируем наличие обязательных полей для RAG2
                        if "article" not in article:
                            article["article"] = "Статья (без названия)"
                        if "relevance_reasoning" not in article:
                            article["relevance_reasoning"] = "Релевантность не указана"
                        if "exact_match" not in article:
                            article["exact_match"] = False

                        # 🔥 КРИТИЧЕСКОЕ ИЗМЕНЕНИЕ: ГАРАНТИРУЕМ наличие поля 'law',
                        # беря его из исходной главы
                        # chapter['law'] содержит имя файла, например "ФЗ-127 О
                        # банкротстве.docx"
                        original_law_name = chapter.get('law', 'Unknown')

                        # 🔥 ВСЕГДА перезаписываем поле law из главы, даже если ИИ указал другое
                        if article.get('law') and article['law'] != original_law_name:
                            print(f"        ⚠️  Статья {article.get('article', 'Unknown')}: "
                                  f"ИИ указал закон '{article['law']}', но заменён на '{original_law_name}'")

                        # Гарантированно устанавливаем правильное название закона
                        article['law'] = original_law_name

                        # 🔥 Дополнительная проверка: логируем, если поле было пустым
                        if not article.get('law'):
                            print(f"        🚨 КРИТИЧЕСКОЕ ПРЕДУПРЕЖДЕНИЕ: "
                                  f"Статья '{article.get('article', 'Unknown')}' без названия закона!")

                        valid_articles.append(article)
                        print(f"        📄 Подготовлена статья для RAG2: "
                              f"{article.get('article', 'Unknown')} (закон: {article['law']})")

                result["articles_found"] = valid_articles
                print(
                    f"        ✅ Успешно подготовлено для RAG2: {len(valid_articles)} статей")
                return result

            else:
                print(f"        ❌ Не удалось найти JSON в ответе")
                return {
                    "articles_found": [],
                    "analysis_quality": "parse_error",
                    "error": "Не удалось распарсить ответ ИИ"
                }

        except json.JSONDecodeError as e:
            print(f"        ❌ Ошибка парсинга JSON: {e}")
            
            # 🔥 ДОБАВЛЕНО: Расширенное сохранение ответа для отладки
            debug_dir = self.project_root / "debug" / "parse_errors"
            debug_dir.mkdir(parents=True, exist_ok=True)
            timestamp = int(time.time())
            law_name = chapter.get('law', 'unknown').replace('.', '_')
            chapter_title = chapter.get('title', 'unknown').replace('/', '_').replace('\\', '_')
            safe_chapter_title = re.sub(r'[^\w\s-]', '', chapter_title)[:50]
            filename = f"{timestamp}_{law_name}_{safe_chapter_title}_outer.json"
            filepath = debug_dir / filename
            
            debug_data = {
                "timestamp": timestamp,
                "chapter": {
                    "title": chapter.get('title', 'unknown'),
                    "law": chapter.get('law', 'unknown'),
                    "content_length": len(original_response)
                },
                "error": {
                    "type": "json_decode_error",
                    "message": str(e),
                    "position": e.pos if hasattr(e, 'pos') else None,
                    "line": e.lineno if hasattr(e, 'lineno') else None,
                    "col": e.colno if hasattr(e, 'colno') else None
                },
                "response_preview": original_response[:2000],
                "json_str_preview": json_str[:2000] if 'json_str' in locals() and json_str else None,
                "full_response_path": str(debug_dir / f"{timestamp}_{law_name}_{safe_chapter_title}_full.txt")
            }
            
            # Сохраняем структурированные данные
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(debug_data, f, ensure_ascii=False, indent=2)
            
            # Сохраняем полный ответ в отдельный файл
            full_response_file = debug_dir / f"{timestamp}_{law_name}_{safe_chapter_title}_full.txt"
            with open(full_response_file, 'w', encoding='utf-8') as f:
                f.write("=== ПОЛНЫЙ ОТВЕТ ИИ ===\n")
                f.write(original_response)
                f.write("\n\n=== JSON_STR ДЛЯ ПАРСИНГА ===\n")
                f.write(json_str if 'json_str' in locals() and json_str else "NOT EXTRACTED")
            
            print(f"        🔍 Проблемный ответ сохранен в: {filepath}")
            print(f"        📄 Полный ответ сохранен в: {full_response_file}")

            # ДОБАВЛЕНО: Расширенная диагностика ошибки
            print(f"        🔍 Диагностика ошибки JSON:")
            print(f"            Позиция ошибки: {e.pos}")
            print(f"            Строка ошибки: {e.lineno}, колонка: {e.colno}")

            # Показываем контекст вокруг ошибки
            if hasattr(e, 'pos') and e.pos and len(original_response) > e.pos:
                start = max(0, e.pos - 100)
                end = min(len(original_response), e.pos + 100)
                context = original_response[start:end]
                print(f"            Контекст ошибки: ...{context}...")

            # Попытка восстановить JSON при ошибке
            try:
                if 'json_str' in locals() and json_str:
                    repaired_json = self._repair_json_string(json_str)
                    if repaired_json:
                        print(f"        🔧 JSON восстановлен после ошибки")

                        # Применяем ту же логику подготовки статей для RAG2
                        result = json.loads(repaired_json)
                        if "articles_found" not in result:
                            result["articles_found"] = []
                        if "analysis_quality" not in result:
                            result["analysis_quality"] = "unknown"

                        valid_articles = []
                        for article in result.get("articles_found", []):
                            if isinstance(article, dict) and article.get('content'):
                                if "article" not in article:
                                    article["article"] = "Статья (без названия)"
                                if "relevance_reasoning" not in article:
                                    article["relevance_reasoning"] = "Релевантность не указана"
                                if "exact_match" not in article:
                                    article["exact_match"] = False
                                if "law" not in article:
                                    article["law"] = chapter.get('law', 'Unknown')

                                valid_articles.append(article)
                                print(f"        📄 Восстановлена статья для RAG2: {article.get('article', 'Unknown')}")

                        result["articles_found"] = valid_articles
                        return result
            except Exception as repair_error:
                print(f"        ❌ Не удалось восстановить JSON: {repair_error}")

            return {
                "articles_found": [],
                "analysis_quality": "json_error",
                "error": f"Ошибка JSON: {str(e)}"
            }
        except Exception as e:
            print(f"        ❌ Неожиданная ошибка: {e}")

            # 🔥 В КОНЦЕ МЕТОДА (перед последним return) ДОБАВИТЬ:
            # Сохраняем проблемный ответ для отладки
            debug_dir = self.project_root / "debug" / "parse_errors"
            debug_dir.mkdir(parents=True, exist_ok=True)
            timestamp = int(time.time())
            law_name = chapter.get('law', 'unknown').replace('.', '_')
            chapter_title = chapter.get('title', 'unknown').replace('/', '_').replace('\\', '_')
            safe_chapter_title = re.sub(r'[^\w\s-]', '', chapter_title)[:50]
            filename = f"{timestamp}_{law_name}_{safe_chapter_title}_exception.txt"
            filepath = debug_dir / filename
            
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(f"\n--- BROKEN RESPONSE [{time.ctime()}] ---\n")
                    f.write(f"Глава: {chapter.get('title', 'Unknown')}\n")
                    f.write(f"Закон: {chapter.get('law', 'Unknown')}\n")
                    f.write(f"Ошибка: {str(e)}\n")
                    f.write(f"Ответ (первые 5000 символов):\n{cleaned_response[:5000]}\n")
                    f.write("--- КОНЕЦ ---\n")
                print(f"        🔍 Проблемный ответ сохранен в: {filepath}")
            except Exception as debug_e:
                print(f"        ⚠️ Не удалось сохранить отладочный ответ: {debug_e}")

            # 🔥 ВОЗВРАЩАЕМ ЗАГЛУШКУ ВМЕСТО None
            return {
                "articles_found": [],
                "analysis_quality": "parse_error",
                "error": f"Не удалось распарсить ответ ИИ: {str(e)}",
                "reasoning": "Ошибка синтаксиса JSON в ответе модели"
            }    
    
    def _clean_json_string(self, json_str: str) -> str:
        """Очищает строку JSON от распространенных проблем с сохранением содержания статей"""
        if not json_str:
            return json_str
        
        # Удаляем лишние пробелы в начале/конце
        json_str = json_str.strip()
        
        # Удаляем BOM символы если есть
        json_str = json_str.replace('\ufeff', '')
        
        # 🔥 ВАЖНОЕ ИЗМЕНЕНИЕ: Сохраняем содержание статей до обработки
        # Извлекаем все поля content для защиты от повреждения
        content_matches = list(re.finditer(r'"content"\s*:\s*"([^"]*)"', json_str, re.DOTALL))
        content_preserve = {}
        
        for i, match in enumerate(content_matches):
            original_content = match.group(1)
            # Временно заменяем содержание на плейсхолдер
            placeholder = f"__CONTENT_PLACEHOLDER_{i}__"
            json_str = json_str.replace(f'"content": "{original_content}"', f'"content": "{placeholder}"')
            content_preserve[placeholder] = original_content
        
        # Теперь безопасно обрабатываем JSON структуру
        # Заменяем одиночные кавычки на двойные (только для ключей и значений, не для плейсхолдеров)
        json_str = re.sub(r'(?<!\\)\'', '"', json_str)
        
        # Удаляем лишние запятые в конце объектов и массивов
        json_str = re.sub(r',\s*}', '}', json_str)
        json_str = re.sub(r',\s*]', ']', json_str)
        
        # Удаляем комментарии (если есть)
        json_str = re.sub(r'//.*?\n', '', json_str)
        json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)
        
        # Исправляем пропущенные запятые между элементами массива
        json_str = re.sub(r'("\s*")\s*"', r'\1, "', json_str)
        json_str = re.sub(r'(true|false|null)\s*"', r'\1, "', json_str)
        json_str = re.sub(r'(\d+)\s*"', r'\1, "', json_str)
        json_str = re.sub(r'}\s*{', '}, {', json_str)
        
        # 🔥 ВОССТАНАВЛИВАЕМ содержание статей после обработки
        for placeholder, original_content in content_preserve.items():
            # Экранируем кавычки в содержании для безопасной вставки в JSON
            safe_content = original_content.replace('"', '\\"').replace('\n', '\\n')
            json_str = json_str.replace(f'"{placeholder}"', f'"{safe_content}"')
        
        return json_str        
        # 🔥 ДОБАВИТЬ ЭТО В КОНЕЦ МЕТОДА:

        # 6. Удаляем Markdown-обертки если они есть
        json_str = re.sub(r'^```(?:json)?\s*', '', json_str, flags=re.IGNORECASE)
        json_str = re.sub(r'\s*```$', '', json_str)

        # 7. Заменяем переносы строк внутри JSON на пробелы (но сохраняем \n в строках)
        # Это решает проблему Expecting ',' delimiter
        lines = json_str.split('\n')
        cleaned_lines = []
        in_string = False
        for line in lines:
            # Простой детектор нахождения внутри строки
            if not in_string and '"' in line:
                # Считаем кавычки
                quote_count = line.count('"') - line.count('\\"')
                if quote_count % 2 != 0:
                    in_string = not in_string
            cleaned_lines.append(line.strip() if not in_string else line)

        # Объединяем: вне строк склеиваем без переносов, внутри строк сохраняем переносы
        result = []
        for line in cleaned_lines:
            if in_string:
                result.append(line + ' ')
            else:
                result.append(line)

        json_str = ' '.join(result)

        # 8. Исправляем распространенные паттерны ошибок
        # Паттерн: {"key": "value" "key2": "value2"} - пропущена запятая
        json_str = re.sub(r'("\s*"\s*:\s*"[^"]*")\s*"', r'\1, "', json_str)

        return json_str
    
    
    def _repair_json_string(self, json_str: str) -> str:
        """Пытается восстановить поврежденный JSON с улучшенной обработкой длинного содержания"""
        if not json_str:
            return None
        
        # 🔥 ДОБАВЛЕНО: Защита содержания статей перед восстановлением
        content_matches = list(re.finditer(r'"content"\s*:\s*"([^"]*)"', json_str, re.DOTALL))
        content_preserve = {}
        
        for i, match in enumerate(content_matches):
            original_content = match.group(1)
            placeholder = f"__CONTENT_PLACEHOLDER_{i}__"
            json_str = json_str.replace(f'"content": "{original_content}"', f'"content": "{placeholder}"')
            content_preserve[placeholder] = original_content
        
        try:
            # Стандартное восстановление структуры
            repaired = json_str.strip()
            repaired = self._fix_missing_commas(repaired)
            repaired = self._fix_unclosed_quotes(repaired)
            repaired = self._fix_unclosed_brackets(repaired)
            
            # 🔥 ВОССТАНАВЛИВАЕМ содержание
            for placeholder, original_content in content_preserve.items():
                safe_content = original_content.replace('"', '\\"').replace('\n', '\\n')
                repaired = repaired.replace(f'"{placeholder}"', f'"{safe_content}"')
            
            # Проверяем, что восстановленный JSON валиден
            json.loads(repaired)
            return repaired
        except Exception:
            pass
        
        # Если стандартное восстановление не помогло, пробуем извлечь JSON по балансу скобок
        try:
            brace_count = 0
            square_count = 0
            start_index = -1
            end_index = -1
            
            for i, char in enumerate(json_str):
                if char == '{':
                    if brace_count == 0:
                        start_index = i
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0 and start_index != -1:
                        end_index = i + 1
                        break
                elif char == '[':
                    square_count += 1
                elif char == ']':
                    square_count -= 1
            
            if start_index != -1 and end_index != -1:
                extracted = json_str[start_index:end_index]
                # Рекурсивно чистим извлеченный JSON
                cleaned_extracted = self._clean_json_string(extracted)
                
                # 🔥 ВОССТАНАВЛИВАЕМ содержание в извлеченном JSON
                for placeholder, original_content in content_preserve.items():
                    safe_content = original_content.replace('"', '\\"').replace('\n', '\\n')
                    cleaned_extracted = cleaned_extracted.replace(f'"{placeholder}"', f'"{safe_content}"')
                
                json.loads(cleaned_extracted)
                return cleaned_extracted
        except:
            pass
            # 6. Удаляем Markdown-обертки если они есть
        json_str = re.sub(r'^```(?:json)?\s*', '', json_str, flags=re.IGNORECASE)
        json_str = re.sub(r'\s*```$', '', json_str)

        # 7. Заменяем переносы строк внутри JSON на пробелы (но сохраняем \n в строках)
        # Это решает проблему Expecting ',' delimiter
        lines = json_str.split('\n')
        cleaned_lines = []
        in_string = False
        for line in lines:
            # Простой детектор нахождения внутри строки
            if not in_string and '"' in line:
                # Считаем кавычки
                quote_count = line.count('"') - line.count('\\"')
                if quote_count % 2 != 0:
                    in_string = not in_string
            cleaned_lines.append(line.strip() if not in_string else line)

        # Объединяем: вне строк склеиваем без переносов, внутри строк сохраняем переносы
        result = []
        for line in cleaned_lines:
            if in_string:
                result.append(line + ' ')
            else:
                result.append(line)

        json_str = ' '.join(result)

        # 8. Исправляем распространенные паттерны ошибок
        # Паттерн: {"key": "value" "key2": "value2"} - пропущена запятая
        json_str = re.sub(r'("\s*"\s*:\s*"[^"]*")\s*"', r'\1, "', json_str)

        return json_str
        
    # 🔥 ДОБАВИТЬ ЭТО ПОСЛЕ СУЩЕСТВУЮЩЕГО КОДА:
        try:
            # Стратегия 3: Агрессивное восстановление - находим самый внешний JSON объект
            # Ищем первый { и последний } и пытаемся извлечь
            first_brace = json_str.find('{')
            last_brace = json_str.rfind('}')
            
            if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
                extracted = json_str[first_brace:last_brace + 1]
                # Очищаем извлеченный фрагмент
                extracted = self._clean_json_string(extracted)
                
                # 🔥 УБИРАЕМ ПЕРЕНОСЫ СТРОК МЕЖДУ ЭЛЕМЕНТАМИ
                # Заменяем переносы, которые не внутри строк, на запятые или удаляем
                lines = extracted.split('\n')
                fixed_lines = []
                for line in lines:
                    line = line.strip()
                    if line and not line.endswith(',') and not line.endswith('{') and not line.endswith('['):
                        line = line + ','
                    fixed_lines.append(line)
                
                extracted = ' '.join(fixed_lines)
                
                # Исправляем лишние запятые
                extracted = re.sub(r',\s*}', '}', extracted)
                extracted = re.sub(r',\s*]', ']', extracted)
                
                # Восстанавливаем защищенный контент
                for placeholder, original_content in content_preserve.items():
                    safe_content = original_content.replace('"', '\\"').replace('\n', '\\n')
                    extracted = extracted.replace(f'"{placeholder}"', f'"{safe_content}"')
                
                json.loads(extracted)
                return extracted
        except Exception:
            pass

        # 🔥 ЕСЛИ НИЧЕГО НЕ ПОМОГЛО, ПРОБУЕМ ВЕРНУТЬ ПРОСТУЮ ЗАГЛУШКУ
        try:
            # Возвращаем валидный JSON с пустыми данными
            return '{"articles_found": [], "analysis_quality": "error"}'
        except:
            return None    
    
    def _fix_missing_commas(self, json_str: str) -> str:
        """Исправляет пропущенные запятые между элементами JSON"""
        repaired = json_str
        
        # Паттерны для поиска мест, где должны быть запятые
        patterns = [
            # Между закрывающей кавычкой и открывающей кавычкой
            (r'"\s*"', '", "'),
            # После строки перед открывающей фигурной скобкой
            (r'"\s*{', '", {'),
            # После строки перед открывающей квадратной скобкой  
            (r'"\s*\[', '", ['),
            # После числа перед открывающей кавычкой
            (r'(\d+)\s*"', r'\1, "'),
            # После boolean/null перед открывающей кавычкой
            (r'(true|false|null)\s*"', r'\1, "'),
            # Между объектами в массиве
            (r'}\s*{', '}, {'),
        ]
        
        for pattern, replacement in patterns:
            repaired = re.sub(pattern, replacement, repaired)
        
        return repaired

    def _fix_unclosed_quotes(self, json_str: str) -> str:
        """Исправляет незакрытые кавычки"""
        lines = json_str.split('\n')
        repaired_lines = []
        
        for line in lines:
            # Считаем кавычки в строке
            quote_count = line.count('"')
            if quote_count % 2 != 0:  # Нечетное количество кавычек
                # Добавляем закрывающую кавычку в конце
                line = line + '"'
            repaired_lines.append(line)
        
        return '\n'.join(repaired_lines)

    def _fix_unclosed_brackets(self, json_str: str) -> str:
        """Исправляет незакрытые скобки"""
        brace_count = json_str.count('{') - json_str.count('}')
        square_count = json_str.count('[') - json_str.count(']')
        
        repaired = json_str
        
        # Добавляем недостающие закрывающие скобки
        if brace_count > 0:
            repaired += '}' * brace_count
        if square_count > 0:
            repaired += ']' * square_count
        
        return repaired

    def _is_irrelevant_article(self, content: str, reasoning: str) -> bool:
        """Проверяет, является ли статья нерелевантной - СИЛЬНО УПРОЩЕНА"""
        
        # ИЗМЕНЕНИЕ: Убираем агрессивную фильтрацию, пусть RAG2 сам решает
        # Оставляем только очевидные случаи нерелевантности
        
        content_lower = content.lower()
        reasoning_lower = reasoning.lower()
        
        # Только явно нерелевантные случаи
        irrelevant_indicators = [
            "коап", "административных правонарушениях", "административный кодекс"
        ]
        
        # Если статья явно из другого кодекса
        for indicator in irrelevant_indicators:
            if indicator in content_lower or indicator in reasoning_lower:
                print(f"   ⚠️  Пропущена явно нерелевантная статья: {indicator}")
                return True
        
        return False
    
    def _save_final_articles_to_json(self, articles: List[Dict], question: str, category: str):
        """Сохраняет валидированные статьи в структурированный JSON файл"""
        data = {
            "metadata": {
                "question": question,
                "category": category,
                "timestamp": time.time(),
                "count": len(articles)
            },
            "articles": articles
        }
        
        try:
            with open(self.articles_file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"   ✅ Статьи успешно сохранены в JSON: {self.articles_file_path}")
        except Exception as e:
            print(f"   ❌ Ошибка сохранения статей в JSON: {e}")
            # Fallback: если не удалось сохранить JSON, попробуем хотя бы лог обновить
            self.search_log["error_saving_json"] = str(e)