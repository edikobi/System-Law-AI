# src/tools/structure_analyzer.py
import os
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches

class DocumentStructureAnalyzer:
    """Анализатор структуры документов и чанков"""
    
    def __init__(self, project_root=None):
        if project_root:
            self.project_root = Path(project_root)
        else:
            self.project_root = Path(__file__).parent.parent.parent
        
        self.laws_dir = self.project_root / "laws"
        self.output_dir = self.project_root / "structure_analysis"
        self.output_dir.mkdir(exist_ok=True)
        
        # Импортируем необходимые компоненты
        import sys
        sys.path.insert(0, str(self.project_root))
        
        from src.core.law_navigator import LawNavigator
        from src.core.sequential_searcher import IntelligentChunkManager
        
        self.navigator = LawNavigator()
        self.chunk_manager = IntelligentChunkManager()
        
        print(f"📁 Анализ структуры документов в: {self.laws_dir}")
    
    def analyze_all_documents(self):
        """Анализирует структуру всех документов"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        txt_file = self.output_dir / f"document_structures_{timestamp}.txt"
        docx_file = self.output_dir / f"document_structures_{timestamp}.docx"
        
        print(f"📊 Начинаем анализ структуры документов...")
        
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write("ПОЛНЫЙ АНАЛИЗ СТРУКТУРЫ ДОКУМЕНТОВ\n")
            f.write("=" * 80 + "\n")
            f.write(f"Время анализа: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Всего документов: {len(self.navigator.law_structures)}\n\n")
            
            for law_name, law_data in self.navigator.law_structures.items():
                f.write(f"\n{'='*60}\n")
                f.write(f"ДОКУМЕНТ: {law_name}\n")
                f.write(f"ТИП: {law_data.get('document_type', 'unknown')}\n")
                f.write(f"СТРУКТУРНЫХ ЕДИНИЦ: {len(law_data['structure'])}\n")
                f.write(f"{'='*60}\n\n")
                
                # Анализируем структуру документа и чанки
                self._analyze_document_chunks(f, law_name, law_data)
        
        # Создаем DOCX версию
        self._create_docx_report(docx_file, txt_file)
        
        print(f"✅ Анализ завершен!")
        print(f"📄 TXT отчет: {txt_file}")
        print(f"📄 DOCX отчет: {docx_file}")
        
        return str(txt_file), str(docx_file)
    
    def _analyze_document_chunks(self, file, law_name: str, law_data: dict):
        """Анализирует чанки и подчанки документа"""
        structures = law_data['structure']
        doc_type = law_data.get('document_type', 'unknown')
        
        file.write("СТРУКТУРА ДОКУМЕНТА (ЧАНКИ):\n")
        file.write("-" * 50 + "\n")
        
        for i, structure in enumerate(structures, 1):
            structure_type = structure.get('type', 'unknown')
            title = structure.get('title', 'Без названия')
            level = structure.get('level', 0)
            size = len(structure.get('content', ''))
            
            # Определяем иконку для типа структуры
            type_icons = {
                'section': '📚', 'chapter': '📖', 'subsection': '📑',
                'major_article': '📄', 'article': '📃', 'preamble': '📜',
                'plenum_chapter': '📘', 'large_block': '📦'
            }
            icon = type_icons.get(structure_type, '🔸')
            
            file.write(f"{i:2d}. {icon} [{structure_type}] Уровень {level}: {title}\n")
            file.write(f"     Размер: {size:,} символов | Строк: {structure.get('content', '').count(chr(10)) + 1}\n")
            
            # Анализируем подчанки для этого чанка
            self._analyze_subchunks(file, structure)
            
            file.write("\n")
    
    def _analyze_subchunks(self, file, structure: dict):
        """Анализирует подчанки для данного чанка"""
        content = structure.get('content', '')
        title = structure.get('title', 'Без названия')
        
        # Подготавливаем чанки для анализа
        chunks = self.chunk_manager.prepare_chapter_analysis(structure)
        
        if len(chunks) == 1:
            # Один чанк - не разбивается на подчанки
            file.write(f"     └─ 📄 Один чанк (не разбивается)\n")
            return
        
        # Множественные подчанки
        file.write(f"     └─ 🪚 Разбит на {len(chunks)} подчанков:\n")
        
        for j, chunk in enumerate(chunks, 1):
            chunk_type = chunk.get('type', 'unknown')
            chunk_title = chunk.get('title', 'Без названия')
            chunk_size = len(chunk.get('content', ''))
            chunk_index = chunk.get('chunk_index', j-1)
            total_chunks = chunk.get('total_chunks', len(chunks))
            
            file.write(f"        {j:2d}. 📦 [{chunk_type}] {chunk_title}\n")
            file.write(f"            Размер: {chunk_size:,} символов | Чанк {chunk_index + 1}/{total_chunks}\n")
            
            # Анализируем границы подчанка
            self._analyze_chunk_boundaries(file, chunk, indent=16)
    
    def _analyze_chunk_boundaries(self, file, chunk: dict, indent: int = 4):
        """Анализирует границы чанка"""
        content = chunk.get('content', '')
        if not content:
            return
        
        lines = content.split('\n')
        indent_str = ' ' * indent
        
        # Анализируем начало чанка
        if len(lines) > 0:
            first_lines = []
            for i in range(min(3, len(lines))):
                line = lines[i].strip()
                if line and len(line) > 10:  # Не пустая и не слишком короткая
                    first_lines.append(line[:100])
                    if len(first_lines) >= 2:  # Берем первые 2 непустые строки
                        break
            
            if first_lines:
                file.write(f"{indent_str}📖 Начало:\n")
                for line in first_lines:
                    file.write(f"{indent_str}   {line}...\n")
        
        # Анализируем конец чанка  
        if len(lines) > 5:
            last_lines = []
            for i in range(max(0, len(lines)-3), len(lines)):
                line = lines[i].strip()
                if line and len(line) > 10:
                    last_lines.append(line[:100])
                    if len(last_lines) >= 2:
                        break
            
            if last_lines:
                file.write(f"{indent_str}📚 Конец:\n")
                for line in last_lines:
                    file.write(f"{indent_str}   ...{line}\n")
        
        # Ищем структурные маркеры в чанке
        markers_found = []
        for line in lines[:10]:  # Первые 10 строк
            if any(marker in line.upper() for marker in 
                  ['ГЛАВА', 'РАЗДЕЛ', 'СТАТЬЯ', 'ПАРАГРАФ', '§', 'СТ.']):
                clean_line = line.strip()[:80]
                if clean_line and clean_line not in markers_found:
                    markers_found.append(clean_line)
                if len(markers_found) >= 3:  # Ограничиваем количество
                    break
        
        if markers_found:
            file.write(f"{indent_str}🔍 Структурные маркеры:\n")
            for marker in markers_found:
                file.write(f"{indent_str}   • {marker}\n")
    
    def _create_docx_report(self, docx_file: Path, txt_file: Path):
        """Создает DOCX версию отчета"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('Анализ структуры юридических документов', 0)
        doc.add_paragraph(f"Время анализа: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Всего документов: {len(self.navigator.law_structures)}")
        doc.add_paragraph()
        
        # Читаем TXT файл и добавляем в DOCX
        with open(txt_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            if line.startswith('==='):
                # Разделитель
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            elif line.startswith('ДОКУМЕНТ:'):
                # Заголовок документа
                doc.add_heading(line, level=1)
            elif line.startswith('СТРУКТУРА ДОКУМЕНТА'):
                # Подзаголовок
                doc.add_heading(line, level=2)
            elif any(line.startswith(str(i) + '.') for i in range(1, 100)):
                # Элемент списка чанков
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            elif line.startswith('     └─'):
                # Подчанки
                p = doc.add_paragraph(line)
                p.style = 'List Bullet 2'
            elif line.startswith('        '):
                # Элементы подчанков
                p = doc.add_paragraph(line)
                p.style = 'List Bullet 3'
            else:
                # Обычный текст
                doc.add_paragraph(line)
        
        doc.save(docx_file)

    def analyze_specific_document(self, law_name: str):
        """Анализирует структуру конкретного документа"""
        if law_name not in self.navigator.law_structures:
            print(f"❌ Документ '{law_name}' не найден")
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        txt_file = self.output_dir / f"{law_name.replace('.txt', '')}_structure_{timestamp}.txt"
        
        law_data = self.navigator.law_structures[law_name]
        
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(f"ДЕТАЛЬНЫЙ АНАЛИЗ: {law_name}\n")
            f.write("=" * 80 + "\n")
            f.write(f"Тип документа: {law_data.get('document_type', 'unknown')}\n")
            f.write(f"Полный размер: {len(law_data.get('full_content', '')):,} символов\n\n")
            
            self._analyze_document_chunks(f, law_name, law_data)
            
            # Дополнительный анализ: поиск конкретных терминов
            self._analyze_term_occurrence(f, law_name, law_data)
        
        print(f"✅ Анализ '{law_name}' завершен: {txt_file}")
        return str(txt_file)
    
    def _analyze_term_occurrence(self, file, law_name: str, law_data: dict):
        """Анализирует встречаемость ключевых терминов"""
        file.write("\nАНАЛИЗ КЛЮЧЕВЫХ ТЕРМИНОВ:\n")
        file.write("-" * 50 + "\n")
        
        content = law_data.get('full_content', '').lower()
        terms_to_search = [
            'внесудебн', 'банкротств', 'несостоятельност',
            'должник', 'кредитор', 'реструктуризац', 'мировое соглашение'
        ]
        
        for term in terms_to_search:
            count = content.count(term)
            if count > 0:
                file.write(f"🔍 '{term}': встречается {count} раз\n")
                
                # Находим контекст первого упоминания
                first_occurrence = content.find(term)
                if first_occurrence != -1:
                    context_start = max(0, first_occurrence - 100)
                    context_end = min(len(content), first_occurrence + 100)
                    context = content[context_start:context_end].replace('\n', ' ')
                    file.write(f"   Первое упоминание: ...{context}...\n")
            else:
                file.write(f"❌ '{term}': не найден\n")

# Утилита для быстрого запуска
def analyze_documents(project_root=None, specific_document=None):
    """Основная функция для анализа документов"""
    analyzer = DocumentStructureAnalyzer(project_root)
    
    if specific_document:
        return analyzer.analyze_specific_document(specific_document)
    else:
        return analyzer.analyze_all_documents()

if __name__ == "__main__":
    # Пример использования
    txt_path, docx_path = analyze_documents()
    print(f"📊 Отчеты созданы:\n- {txt_path}\n- {docx_path}")